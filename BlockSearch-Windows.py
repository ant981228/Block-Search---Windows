#!/usr/bin/env python3
"""
Advanced Word Document Search and Content Transfer Utility

Provides sophisticated document searching capabilities with support for both clipboard 
operations and direct content transfer between documents, including intelligent handling 
of open Word documents with multiple paste modes.
"""

# Standard library imports
import csv
import ctypes
import docx
import os
import re
import sys
import zipfile
from ctypes import wintypes
from contextlib import contextmanager
from dataclasses import dataclass, field
from datetime import datetime
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Dict, List, Optional, Set, Any, Tuple

# Third-party imports
import keyboard
import pythoncom
import win32com.client

# PyQt imports - grouped by functionality
from PyQt6.QtNetwork import QLocalServer, QLocalSocket

from PyQt6.QtCore import (
    Qt,
    QTimer,
    QSettings,
    QEvent,
    pyqtSignal
)

from PyQt6.QtGui import (
    QColor,
    QPalette,
    QAction,
    QActionGroup,
    QKeyEvent,
    QKeySequence,
    QIcon,
    QShortcut,
    QTextOption
)

from PyQt6.QtWidgets import (
    QApplication,
    QDialog,
    QDialogButtonBox,
    QFileDialog,
    QFrame,
    QHBoxLayout,
    QHeaderView,
    QInputDialog,
    QLabel,
    QLineEdit,
    QListWidget,
    QListWidgetItem,
    QMainWindow,
    QMenu,
    QMenuBar,
    QMessageBox,
    QPushButton,
    QSplitter,
    QStatusBar,
    QSystemTrayIcon,
    QTabWidget,
    QTableWidget,
    QTableWidgetItem,
    QTextEdit,
    QVBoxLayout,
    QWidget,
    QMenu as QTrayMenu, 
    QComboBox,           
    QProgressBar,        
    QCheckBox            
)

@dataclass
class PrefixConfig:
    """
    Represents a prefix configuration with its associated folders.
    
    Attributes:
        prefix: The search prefix (e.g., "th")
        folders: Set of folder paths associated with this prefix
    """
    prefix: str
    folders: Set[str]

@dataclass
class DocumentInfo:
    """
    Enhanced document information container with comprehensive metadata support.
    
    Attributes:
        path: Path object pointing to document location
        name: Document filename
        last_modified: Unix timestamp of last modification
        created_time: Unix timestamp of creation time
        size: File size in bytes
        search_name: Lowercase name for efficient search operations
        relative_path: Path relative to search root for folder display
        original_doc_path: Path to the original source document (if split from larger doc)
        position_in_original: Position/index in the original document
        parent_doc_name: Name of the parent document (if applicable)
        sibling_docs: List of sibling document names in document order
    """
    path: Path
    name: str
    last_modified: float
    created_time: float
    size: int
    search_name: str = ""
    relative_path: str = ""
    original_doc_path: Optional[str] = None
    position_in_original: Optional[int] = None
    parent_doc_name: Optional[str] = None
    sibling_docs: List[str] = field(default_factory=list)

    def __post_init__(self):
        """Initialize derived attributes for optimized search operations."""
        self.search_name = self.name.lower()
        
@dataclass
class ActiveDocument:
    """
    Represents an active Word document with its associated window and cursor state.
    
    Attributes:
        name: Display name of the document
        path: Full path to the document file
        window_index: Index in Word's window collection
        doc_id: Unique document identifier in Word's session
    """
    name: str
    path: str
    window_index: int
    doc_id: str

@dataclass
class PasteMode:
    """Type-safe enumeration of paste modes with clear intent."""
    CURSOR = "cursor"
    END = "end"

@dataclass
class Section:
    """
    Represents a document section with its heading and content.
    """
    title: str
    safe_title: str
    level: int
    content: List[Paragraph]
    start_index: int
    end_index: Optional[int] = None
    parent: Optional['Section'] = None
    children: List['Section'] = field(default_factory=list)
    
    def get_path_components(self) -> List[str]:
        """Get folder path components based on parent hierarchy."""
        if self.parent is None:
            return []
        else:
            # This recursively builds the path by getting parent components first
            return self.parent.get_path_components() + [self.parent.safe_title]

class HelpDialog(QDialog):
    """
    Comprehensive help dialog explaining application functionality.
    """
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Block Search Help")
        self.setMinimumSize(700, 600)
        
        # Main layout
        layout = QVBoxLayout(self)
        
        # Create a tab widget for organized help sections
        tab_widget = QTabWidget()
        layout.addWidget(tab_widget)
        
        # Add tabs for different help sections
        tab_widget.addTab(self._create_general_tab(), "General")
        tab_widget.addTab(self._create_search_tab(), "Searching your Blocks")
        tab_widget.addTab(self._create_document_tab(), "Sending to Doc")
        tab_widget.addTab(self._create_splitter_tab(), "Document Splitter") 
        tab_widget.addTab(self._create_shortcuts_tab(), "Keyboard Shortcuts")
        
        # Add close button
        close_button = QPushButton("Close")
        close_button.clicked.connect(self.accept)
        layout.addWidget(close_button)
    
    def _create_general_tab(self):
        """Create the general help tab."""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        text = """
        <h2>Block Search Utility</h2>
        <p>Block Search helps you quickly send blocks to your speech document.</p>
        
        <h3>Getting Started:</h3>
        <ol>
            <li>Select the folder with your blocks from <b>Search Settings → Select Search Folder</b></li>
            <li>Type in the search box to find documents</li>
            <li>Click on a document to copy its contents</li>
        </ol>
        """
        
        label = QLabel(text)
        label.setWordWrap(True)
        label.setTextFormat(Qt.TextFormat.RichText)
        
        layout.addWidget(label)
        layout.addStretch(1)
        
        return widget
    
    def _create_search_tab(self):
        """Create the search help tab."""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        text = """
        <h2>Searching your Blocks</h2>
        
        <h3>Basic Search:</h3>
        <p>Type search terms in the search box to find matching documents. 
        Results update as you type.</p>
        
        <h3>Using Prefixes:</h3>
        <p>Prefixes let you limit searches to specific folders:</p>
        <ol>
            <li>Configure prefixes in <b>Search Settings → Prefix Configuration → Manage Prefixes</b></li>
            <li>Use format: <code>[prefix] [search terms]</code></li>
            <li>Example: <code>cb 2ac</code> searches for "2ac" only in folders assigned to "cb" prefix (say, condo bad!)</li>
        </ol>
        
        <h3>Sorting Results:</h3>
        <p>Use the <b>Sort</b> menu to order results by:
        <ul>
            <li>Name</li>
            <li>Date Modified</li>
            <li>Date Created</li>
            <li>Size</li>
        </ul>
        Toggle <b>Reverse Order</b> to switch between ascending and descending order.</p>
        """
        
        label = QLabel(text)
        label.setWordWrap(True)
        label.setTextFormat(Qt.TextFormat.RichText)
        
        layout.addWidget(label)
        layout.addStretch(1)
        
        return widget
    
    def _create_document_tab(self):
        """Create the document operations help tab."""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        text = """
        <h2>Sending to Doc</h2>
        
        <h3>Document Content Transfer:</h3>
        <p>When you click on a document in search results, the content is handled based on your settings:</p>
        <ul>
            <li><b>Clipboard Mode:</b> Content is copied to clipboard (default when no target is set)</li>
            <li><b>Target Document Mode:</b> Content is appended to your selected closed document</li>
            <li><b>Active Document Mode:</b> Content is pasted into an open Word document</li>
        </ul>
        
        <h3>Setting a Target Document:</h3>
        <p>Use <b>Send to Closed Doc → Select Destination</b> to choose a document to receive content.
        Your block will be added to the end of this file when you click.</p>
        
        <h3>Using Open Word Documents:</h3>
        <p>The <b>Send to Open Doc</b> menu shows currently open Word documents:
        <ol>
            <li>Select a document from the menu to set it as the active target</li>
            <li>Use <b>Default Paste Mode</b> to control where content is inserted:
                <ul>
                    <li><b>Paste at Cursor:</b> Inserts at current cursor position</li>
                    <li><b>Paste at Document End:</b> Appends to the end of document</li>
                </ul>
            </li>
            <li>Press <b>Ctrl+Enter</b> to use the alternate paste mode</li>
        </ol>
        """
        
        label = QLabel(text)
        label.setWordWrap(True)
        label.setTextFormat(Qt.TextFormat.RichText)
        
        layout.addWidget(label)
        layout.addStretch(1)
        
        return widget
    
    def _create_splitter_tab(self):
        """Create the document splitter help tab."""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        text = """
        <h2>Document Splitter</h2>
        
        <h3>Overview:</h3>
        <p>The Document Splitter allows you to break a large Word document into smaller files based on heading levels.
        This is useful for extracting blocks from a large file.</p>
        
        <h3>How to Access:</h3>
        <p>Open the splitter from <b>Document Tools → Split Document by Headings</b></p>
        
        <h3>Using the Splitter:</h3>
        <ol>
            <li><b>Select Input Document:</b> Choose the Word document you want to split</li>
            <li><b>Choose Template Document (Recommended):</b> Select a document to use as a template for output files. This should be a blank word document created using your version of Verbatim.</li> 
                <ul>
                    <li>If no template is selected, a minimal default template will be created</li>
                </ul>
            </li>
            <li><b>Select Heading Level:</b> Choose which heading level to split at (Heading 1-4)
                <ul>
                    <li>The document will be divided at each heading of the selected level</li>
                    <li>Each heading and its content will become a separate document</li>
                </ul>
            </li>
            <li><b>Choose Output Options:</b>
                <ul>
                    <li><b>Create ZIP Archive:</b> Package all split documents into a single ZIP file</li>
                    <li><b>Individual Files:</b> Save each section as a separate document</li>
                </ul>
            </li>
            <li><b>Select Output Location:</b> Choose where to save the output files</li>
            <li><b>Process Document:</b> Click the button to start the splitting process</li>
        </ol>
        
        <h3>Tips:</h3>
        <ul>
            <li>Results will be best for files that use Verbatim styles</li>
            <li>Heading 1 will target Pockets, Heading 2 Hats, Heading 3 Blocks, and Heading 4 Tags</li>
            <li>Use the ZIP option for easier file management when creating many documents</li>
        </ul>
        """
        
        label = QLabel(text)
        label.setWordWrap(True)
        label.setTextFormat(Qt.TextFormat.RichText)
        
        layout.addWidget(label)
        layout.addStretch(1)
        
        return widget

    def _create_shortcuts_tab(self):
        """Create the keyboard shortcuts help tab."""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        text = """
        <h2>Keyboard Shortcuts</h2>
        
        <h3>Application Shortcuts:</h3>
        <table border="0" cellspacing="8">
            <tr>
                <td><b>Ctrl+Space</b></td>
                <td>Show/activate application (configurable)</td>
            </tr>
            <tr>
                <td><b>Escape</b></td>
                <td>Hide application window</td>
            </tr>
            <tr>
                <td><b>F5</b></td>
                <td>Refresh open documents list</td>
            </tr>
            <tr>
                <td><b>Ctrl+T</b></td>
                <td>Select (closed) target document</td>
            </tr>
            <tr>
                <td><b>Ctrl+Shift+T</b></td>
                <td>Clear (closed) target document</td>
            </tr>
            <tr>
                <td><b>Ctrl+P</b></td>
                <td>Toggle between paste modes (cursor/end)</td>
            </tr>
            <tr>
                <td><b>Ctrl+Shift+P</b></td>
                <td>Toggle including path names in search</td>
            </tr>
            <tr>
                <td><b>Alt+F4, Ctrl+Q</b></td>
                <td>Quit application</td>
            </tr>
        </table>
        
        <h3>Search Result Navigation:</h3>
        <table border="0" cellspacing="8">
            <tr>
                <td><b>Up/Down</b></td>
                <td>Navigate through search results</td>
            </tr>
            <tr>
                <td><b>Enter</b></td>
                <td>Select document (use default paste mode)</td>
            </tr>
            <tr>
                <td><b>Ctrl+Enter</b></td>
                <td>Select document (use alternate paste mode)</td>
            </tr>
            <tr>
                <td><b>Left Arrow</b></td>
                <td>Show document in its original context</td>
            </tr>
            <tr>
                <td><b>Right Arrow</b></td>
                <td>Return from context view to search results</td>
            </tr>
            <tr>
                <td><b>F1</b></td>
                <td>Show this help dialog</td>
            </tr>
        </table>
        """
        
        label = QLabel(text)
        label.setWordWrap(True)
        label.setTextFormat(Qt.TextFormat.RichText)
        
        layout.addWidget(label)
        layout.addStretch(1)
        
        return widget

class DocumentPreviewDialog(QDialog):
    """
    Document preview dialog with simplified HTML-based rendering.
    
    This dialog provides a simple preview of Word documents by
    extracting content and formatting using python-docx and
    displaying it as HTML in a QTextEdit widget.
    """
    
    def __init__(self, document_path, parent=None):
        super().__init__(parent)
        self.document_path = document_path
        self.temp_files = []  # Track temporary files for cleanup
        
        # Define highlight color mapping
        self.highlight_colors = {
            "Yellow": "#FFFF00",
            "Light Blue": "#ADD8E6",
            "Green": "#90EE90"
        }
        self.current_highlight_color = self.highlight_colors["Yellow"]  # Default
        
        # Setup basic UI
        self.setWindowTitle(f"Preview: {Path(document_path).name}")
        self.resize(900, 700)
        self.setup_ui()
        
        # Load document with progress indication
        self.progress_bar.show()
        QTimer.singleShot(100, self.load_document)
    
    def setup_ui(self):
        """Set up the user interface with flexible layout."""
        layout = QVBoxLayout(self)
        
        # Document info header
        info_layout = QHBoxLayout()
        self.doc_info_label = QLabel(f"Loading document...")
        info_layout.addWidget(self.doc_info_label)
        
        # Highlight color selector
        highlight_label = QLabel("Highlight Color:")
        info_layout.addWidget(highlight_label)
        
        self.highlight_color_combo = QComboBox()
        self.highlight_color_combo.addItems(["Yellow", "Light Blue", "Green"])
        self.highlight_color_combo.setCurrentIndex(0)  # Default to yellow
        self.highlight_color_combo.currentIndexChanged.connect(self.change_highlight_color)
        info_layout.addWidget(self.highlight_color_combo)
        
        layout.addLayout(info_layout)
        
        # Progress bar (initially hidden)
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 0)  # Indeterminate
        self.progress_bar.hide()
        layout.addWidget(self.progress_bar)
        
        # Text viewer
        self.text_viewer = QTextEdit()
        self.text_viewer.setReadOnly(True)
        self.text_viewer.setWordWrapMode(QTextOption.WrapMode.WordWrap)
        layout.addWidget(self.text_viewer)
        
        # Button row
        btn_layout = QHBoxLayout()
        
        # Copy buttons
        self.copy_btn = QPushButton("Copy to Clipboard")
        self.copy_btn.clicked.connect(self.copy_to_clipboard)
        btn_layout.addWidget(self.copy_btn)
        
        # Selection copy buttons  
        self.copy_selection_btn = QPushButton("Copy Selection")
        self.copy_selection_btn.clicked.connect(self.copy_selection)
        btn_layout.addWidget(self.copy_selection_btn)
        
        btn_layout.addStretch()
        
        # Close button
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(self.accept)
        btn_layout.addWidget(close_btn)
        
        layout.addLayout(btn_layout)
    
    def load_document(self):
        """Load document with HTML-based renderer."""
        try:
            self.load_with_fallback()
            self.progress_bar.hide()
        except Exception as e:
            self.doc_info_label.setText(f"Error loading document: {str(e)}")
            print(f"Document preview error: {e}")
            self.progress_bar.hide()
    
    def change_highlight_color(self, index):
        """Update the highlight color and refresh the view."""
        color_name = self.highlight_color_combo.currentText()
        self.current_highlight_color = self.highlight_colors[color_name]
        
        # Refresh the view
        self.load_with_fallback()
    
    def load_with_fallback(self):
        """Load document with improved python-docx based renderer."""
        self.doc_info_label.setText("Loading document...")
        
        try:
            import docx
            document = docx.Document(self.document_path)
            
            # Create HTML content from document with improved formatting
            html_content = ["<html><body style='font-family: Arial, sans-serif; margin: 20px;'>"]
            
            # Process paragraphs with style information
            for para in document.paragraphs:
                # Skip empty paragraphs but preserve as space
                if not para.text.strip():
                    html_content.append("<p>&nbsp;</p>")
                    continue
                    
                # Extract paragraph style information
                para_styles = self.extract_paragraph_style(para)
                style_attr = f" style=\"{para_styles}\"" if para_styles else ""
                
                # Handle different paragraph types based on style
                if para.style and para.style.name.startswith('Heading '):
                    try:
                        level = int(para.style.name.split()[-1])
                        if 1 <= level <= 6:  # Valid HTML heading levels
                            html_content.append(f"<h{level}{style_attr}>{self.format_runs_as_html(para.runs)}</h{level}>")
                        else:
                            # Fallback for unusual heading levels
                            html_content.append(f"<p{style_attr}><strong>{self.format_runs_as_html(para.runs)}</strong></p>")
                    except (ValueError, IndexError):
                        # Fallback if heading level can't be determined
                        html_content.append(f"<p{style_attr}><strong>{self.format_runs_as_html(para.runs)}</strong></p>")
                else:
                    html_content.append(f"<p{style_attr}>{self.format_runs_as_html(para.runs)}</p>")
            
            html_content.append("</body></html>")
            
            # Set the HTML content to the text viewer
            self.text_viewer.setHtml("".join(html_content))
            self.doc_info_label.setText(f"{Path(self.document_path).name}")
            
        except Exception as e:
            print(f"Basic fallback renderer failed: {e}")
            self.text_viewer.setPlainText(f"Error loading document: {e}\n\nDocument path: {self.document_path}")
            self.doc_info_label.setText(f"Error: {str(e)}")
    
    def extract_paragraph_style(self, para):
        """Extract paragraph styling information."""
        styles = []
        
        # Handle alignment
        if hasattr(para, 'paragraph_format') and para.paragraph_format:
            # Alignment
            if hasattr(para.paragraph_format, 'alignment'):
                alignment = para.paragraph_format.alignment
                if alignment == 1:  # WD_ALIGN_LEFT
                    styles.append("text-align: left;")
                elif alignment == 2:  # WD_ALIGN_CENTER
                    styles.append("text-align: center;")
                elif alignment == 3:  # WD_ALIGN_RIGHT
                    styles.append("text-align: right;")
                elif alignment == 4:  # WD_ALIGN_JUSTIFY
                    styles.append("text-align: justify;")
            
            # Indentation
            if hasattr(para.paragraph_format, 'left_indent') and para.paragraph_format.left_indent:
                try:
                    indent_pt = para.paragraph_format.left_indent.pt
                    styles.append(f"padding-left: {indent_pt}pt;")
                except:
                    pass
            
            # Line spacing
            if hasattr(para.paragraph_format, 'line_spacing'):
                try:
                    if para.paragraph_format.line_spacing:
                        styles.append(f"line-height: {para.paragraph_format.line_spacing:.1f};")
                except:
                    pass
            
            # Spacing before/after
            if hasattr(para.paragraph_format, 'space_before') and para.paragraph_format.space_before:
                try:
                    space_pt = para.paragraph_format.space_before.pt
                    styles.append(f"margin-top: {space_pt}pt;")
                except:
                    pass
                    
            if hasattr(para.paragraph_format, 'space_after') and para.paragraph_format.space_after:
                try:
                    space_pt = para.paragraph_format.space_after.pt
                    styles.append(f"margin-bottom: {space_pt}pt;")
                except:
                    pass
        
        # Check for specific style types from the style name
        if para.style:
            style_name = para.style.name.lower()
            
            # Common style identification patterns
            if 'quote' in style_name or 'block' in style_name:
                styles.append("border-left: 3px solid #ccc; padding-left: 10px; font-style: italic;")
            elif 'list' in style_name:
                styles.append("margin-left: 20px;")
        
        return " ".join(styles)
    
    def format_runs_as_html(self, runs):
        """Convert paragraph runs to HTML with enhanced formatting."""
        result = []
        for run in runs:
            # Handle empty runs
            if not run.text:
                continue
                
            # Escape HTML special characters
            text = run.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # Build style attributes
            styles = []
            
            # Basic formatting
            if run.bold:
                styles.append("font-weight: bold;")
            if run.italic:
                styles.append("font-style: italic;")
            if run.underline:
                styles.append("text-decoration: underline;")
            
            # Handle font properties
            if hasattr(run, 'font'):
                # Font name/family
                if run.font.name:
                    styles.append(f"font-family: '{run.font.name}', Arial, sans-serif;")
                
                # Font size
                if run.font.size:
                    try:
                        # Convert half-points to points if necessary
                        pt_size = run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size / 2
                        styles.append(f"font-size: {pt_size}pt;")
                    except:
                        pass
                
                # Font color
                if hasattr(run.font, 'color') and run.font.color and run.font.color.rgb:
                    color = run.font.color.rgb
                    # Make sure we have a proper color format
                    if isinstance(color, str):
                        color = color.replace('0x', '')
                        if len(color) == 6:
                            styles.append(f"color: #{color};")
                
                # Highlight - use our configurable highlight color
                if hasattr(run.font, 'highlight_color') and run.font.highlight_color:
                    styles.append(f"background-color: {self.current_highlight_color};")
                
                # All caps
                if hasattr(run.font, 'all_caps') and run.font.all_caps:
                    styles.append("text-transform: uppercase;")
                
                # Small caps
                if hasattr(run.font, 'small_caps') and run.font.small_caps:
                    styles.append("font-variant: small-caps;")
            
                # Handle special cases based on style names
                if hasattr(run, 'style') and run.style and run.style.name:
                    style_name = run.style.name.lower()
                    
                    # Debug the style name - can be removed later
                    print(f"Style name: {style_name}")
                    
                    # Look for cite or bold anywhere in the style name
                    if 'cite' in style_name:
                        styles.append("font-weight: bold;")
                    elif 'bold' in style_name:
                        styles.append("font-weight: bold;")
                    elif 'underline' in style_name:
                        styles.append("text-decoration: underline;")
                    elif 'emphasis' in style_name:
                        styles.append("font-weight: bold; text-decoration: underline;")
            
            # Create span with styles or just text if no styles
            if styles:
                result.append(f"<span style=\"{' '.join(styles)}\">{text}</span>")
            else:
                result.append(text)
        
        return "".join(result)
    
    def copy_to_clipboard(self):
        """Copy document content to clipboard."""
        try:
            # Copy from text viewer
            self.text_viewer.selectAll()
            self.text_viewer.copy()
            self.text_viewer.moveCursor(self.text_viewer.textCursor().Start)
            self.doc_info_label.setText("Copied document text to clipboard")
        except Exception as e:
            self.doc_info_label.setText(f"Error copying to clipboard: {str(e)}")
    
    def copy_selection(self):
        """Copy selected content to clipboard."""
        try:
            # Get text viewer selection
            cursor = self.text_viewer.textCursor()
            if cursor.hasSelection():
                self.text_viewer.copy()
                self.doc_info_label.setText("Copied selection to clipboard")
            else:
                self.doc_info_label.setText("No text selected")
        except Exception as e:
            self.doc_info_label.setText(f"Error copying selection: {str(e)}")
    
    def closeEvent(self, event):
        """Clean up resources when dialog is closed."""
        # Clean up temporary files
        for temp_file in self.temp_files:
            try:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
            except Exception as e:
                print(f"Error removing temp file {temp_file}: {e}")
        
        super().closeEvent(event)

class StyleProcessor:
    """
    Processes and manages document styles, handling paragraph styles.
    """
    
    def __init__(self, doc: Document):
        self.doc = doc
        self.heading_levels: Dict[str, int] = {}
        self._process_styles()
    
    def _process_styles(self) -> None:
        for style in self.doc.styles:
            if style.type != WD_STYLE_TYPE.PARAGRAPH:
                continue
                
            # Handle built-in heading styles
            if style.name.startswith('Heading '):
                try:
                    level = int(style.name.split()[-1])
                    self.heading_levels[style.name] = level
                except (ValueError, IndexError):
                    continue
            
            # Handle custom styles based on headings
            elif hasattr(style, 'base_style') and style.base_style:
                base_name = style.base_style.name
                if base_name.startswith('Heading '):
                    try:
                        level = int(base_name.split()[-1])
                        self.heading_levels[style.name] = level
                    except (ValueError, IndexError):
                        continue
    
    def get_heading_level(self, paragraph: Paragraph) -> Optional[int]:
        if not paragraph.style or not paragraph.style.name:
            return None
        return self.heading_levels.get(paragraph.style.name)

class FilenameManager:
    """Handles creation and management of safe filenames."""
    
    @staticmethod
    def sanitize_filename(title: str, max_length: int = 240) -> str:
        # Remove invalid chars
        safe = re.sub(r'[<>:"/\\|?*]', '', title)
        # Replace whitespace with underscore
        safe = re.sub(r'\s+', '_', safe)
        # Remove duplicate dots
        safe = re.sub(r'\.+', '.', safe)
        # Truncate if needed
        safe = safe[:max_length]
        # Remove leading/trailing dots and underscores
        return safe.strip('._')
    
    def ensure_unique(self, filename: str, used_names: Set[str]) -> str:
        base = filename
        counter = 1
        
        while filename in used_names:
            name_parts = base.rsplit('.', 1)
            if len(name_parts) > 1:
                filename = f"{name_parts[0]}_{counter}.{name_parts[1]}"
            else:
                filename = f"{base}_{counter}"
            counter += 1
            
        used_names.add(filename)
        return filename

class DocxSplitter:
    """
    Main class for splitting Word documents by heading level.
    Enhanced to support various output options and cancellation.
    """
    
    def __init__(self, input_path: Path, template_path: Path, status_callback=None, progress_callback=None):
        """
        Initialize the splitter with progress reporting.
        
        Args:
            input_path: Path to document to split
            template_path: Path to template document
            status_callback: Optional callback for status updates
            progress_callback: Optional callback for progress updates (0-100)
        """
        self.input_path = input_path
        self.template_path = template_path
        self.status_callback = status_callback or (lambda msg: None)
        self.progress_callback = progress_callback or (lambda percent: None)
        
        # Verify template exists
        if not template_path.exists():
            raise ValueError(f"Template document not found: {template_path}")
        
        self.doc = docx.Document(input_path)
        self.style_processor = StyleProcessor(self.doc)
        self.filename_manager = FilenameManager()
        self.sections: List[Section] = []
        
        # Add a cancel flag
        self.cancel_requested = False
    
    def cancel(self):
        """Request cancellation of the current operation."""
        print("DocxSplitter.cancel() method called!")  # Debug
        self.cancel_requested = True
        print(f"cancel_requested flag set to {self.cancel_requested}")  # Debug
        self.status_callback("Cancellation requested")
    
    def _clean_document(self, doc: Document, target_level: int) -> Document:
        """
        Clean document by removing higher-level headings and empty headers.
        Preserves the target heading level for splitting.
        
        Args:
            doc: Document to clean
            target_level: The heading level we're splitting on (to preserve)
        """
        paragraphs_to_remove = []
        
        for idx, para in enumerate(doc.paragraphs):
            # Check for cancellation during cleaning
            if idx % 100 == 0 and self.cancel_requested:
                self.status_callback("Operation canceled during document cleaning")
                return doc
                
            level = self.style_processor.get_heading_level(para)
            
            # Check if this is a heading
            if level is not None:
                # Remove if:
                # 1. It's a heading that's higher in hierarchy than our target level
                #    BUT is not our target level itself
                # 2. It's an empty heading (any level, including target level)
                if (level < target_level and level != target_level) or not para.text.strip():
                    paragraphs_to_remove.append(idx)
                    # Mark the next paragraph for removal if it's empty
                    if idx + 1 < len(doc.paragraphs) and not doc.paragraphs[idx + 1].text.strip():
                        paragraphs_to_remove.append(idx + 1)
        
        # Remove paragraphs in reverse order to maintain correct indices
        for idx in sorted(paragraphs_to_remove, reverse=True):
            if idx < len(doc.paragraphs):  # Safety check
                p = doc.paragraphs[idx]._element
                p.getparent().remove(p)
        
        return doc
    
    def parse_sections(self, target_level: int = 3) -> None:
        """Parse document into sections with improved parent detection."""
        self.status_callback(f"Parsing document sections at heading level {target_level}...")
        
        # First collect ALL headings at ALL levels with their positions
        headings_by_level = {level: [] for level in range(1, target_level + 1)}
        all_sections = []
        used_titles = set()
        
        # First pass: Find all headings and create section objects
        print(f"Scanning document for headings up to level {target_level}...")
        for idx, para in enumerate(self.doc.paragraphs):
            level = self.style_processor.get_heading_level(para)
            
            if level is not None and 1 <= level <= target_level and para.text.strip():
                safe_title = self.filename_manager.sanitize_filename(para.text)
                unique_title = self.filename_manager.ensure_unique(safe_title, used_titles)
                
                # Create section object
                section = Section(
                    title=para.text,
                    safe_title=unique_title,
                    level=level,
                    content=[],
                    start_index=idx,
                    end_index=None  # Will set later
                )
                
                # Add to our collections
                headings_by_level[level].append(section)
                all_sections.append(section)
                print(f"Found heading level {level}: '{para.text}' at index {idx}")
        
        # Sort all sections by start index to ensure they're in document order
        all_sections.sort(key=lambda s: s.start_index)
        
        # Set end indices based on the next section of any level
        for i in range(len(all_sections) - 1):
            all_sections[i].end_index = all_sections[i + 1].start_index - 1
        
        # Set end index for the last section
        if all_sections:
            all_sections[-1].end_index = len(self.doc.paragraphs) - 1
        
        # Second pass: Fill in content
        for section in all_sections:
            # Use inclusive range from start to end
            for idx in range(section.start_index, section.end_index + 1):
                if idx < len(self.doc.paragraphs):
                    section.content.append(self.doc.paragraphs[idx])
        
        # Third pass: Establish parent-child relationships
        print("Building parent-child relationships...")
        for section in all_sections:
            if section.level > 1:  # Skip level 1 as they have no parents
                # Find closest preceding section with a lower level
                for idx in range(all_sections.index(section) - 1, -1, -1):
                    potential_parent = all_sections[idx]
                    if potential_parent.level < section.level:
                        section.parent = potential_parent
                        potential_parent.children.append(section)
                        print(f"Set parent for '{section.title}' (level {section.level}) -> '{potential_parent.title}' (level {potential_parent.level})")
                        break
        
        # Fourth pass: Collect just the target level sections
        self.sections = [s for s in all_sections if s.level == target_level]
        self.sections = [s for s in self.sections if any(p.text.strip() for p in s.content)]
        
        # Debug information
        sections_with_parent = sum(1 for s in self.sections if s.parent is not None)
        print(f"Sections with parent: {sections_with_parent}/{len(self.sections)}")
        
        # Print paths for a few sections to verify
        for i, section in enumerate(self.sections[:5]):
            if i >= len(self.sections):
                break
            path = section.get_path_components()
            if path:
                print(f"Section {i+1} path: {'/'.join(path)}/{section.safe_title}")
            else:
                print(f"Section {i+1} has no parent path: {section.safe_title}")
        
        self.status_callback(f"Found {len(self.sections)} sections at level {target_level}")
    
    def process_document(self, output_dir: Path, target_level: int = 3, 
                        create_zip: bool = True, preserve_hierarchy: bool = False) -> Path:
        """
        Process the document and output files according to specified options.
        
        Args:
            output_dir: Output directory for files
            target_level: Heading level to split on (default: 3)
            create_zip: Whether to create a zip archive or individual files
            preserve_hierarchy: Whether to preserve document hierarchy in folder structure
            
        Returns:
            Path: Path to created output (zip file or directory) or None if canceled
        """
        # Make sure we have parsed sections
        if not self.sections:
            self.parse_sections(target_level)
            
        # Check if parsing was canceled
        if self.cancel_requested or not self.sections:
            return None
        
        # Create output directory if it doesn't exist
        output_dir.mkdir(parents=True, exist_ok=True)
        
        if create_zip:
            return self._create_zip_archive(output_dir, preserve_hierarchy)
        else:
            return self._save_individual_files(output_dir, preserve_hierarchy)
    
    def _add_document_metadata(self, doc: Document, section: Section, all_sections: List[Section]) -> None:
        """
        Add hierarchy metadata to document custom properties.
        
        Args:
            doc: Document to add metadata to
            section: Current section being processed
            all_sections: All sections in the document for sibling relationships
        """
        try:
            # Create metadata dictionary
            metadata = {
                "original_doc_path": str(self.input_path),
                "position_in_original": section.start_index,
                "section_level": section.level,
                "section_title": section.title,
            }
            
            # Add parent info if available
            if section.parent:
                metadata["parent_doc_name"] = section.parent.safe_title
            
            # Find all sections with the same parent (or no parent) and at the same level
            siblings = []
            for s in all_sections:
                if s.level == section.level:
                    if (section.parent is None and s.parent is None) or \
                       (section.parent is not None and s.parent is not None and 
                        section.parent.safe_title == s.parent.safe_title):
                        siblings.append(s)
            
            # Sort siblings by position in document
            siblings.sort(key=lambda s: s.start_index)
            
            # Store sibling names (excluding self)
            metadata["sibling_docs"] = [s.safe_title for s in siblings if s.safe_title != section.safe_title]
            
            # Still try to add some core properties for backward compatibility,
            # but keep them short to avoid exceeding limits
            core_props = doc.core_properties
            try:
                # Try to add minimal metadata to the document properties
                core_props.identifier = "Original: " + Path(self.input_path).name
                core_props.category = f"position:{section.start_index}"
                if section.parent:
                    core_props.subject = f"parent:{section.parent.safe_title[:50]}" if section.parent.safe_title else ""
            except Exception as e:
                print(f"Warning: Could not add core properties: {e}")
                
            # Return the metadata - it will be saved to a separate file by the caller
            return metadata
                
        except Exception as e:
            print(f"Error creating metadata: {e}")
            return None
    
    def _create_section_document(self, section: Section) -> Document:
        """
        Create new document from section content using template.
        
        Args:
            section: Section to convert to document
            
        Returns:
            Document: New document containing section content with preserved formatting
        """
        # Create new document from template
        new_doc = docx.Document(self.template_path)
        
        # Remove any existing paragraphs from the template
        for p in new_doc.paragraphs[::-1]:  # Iterate in reverse for stable removal
            p._element.getparent().remove(p._element)
        
        # Add the heading as document title (will be first content)
        heading = new_doc.add_heading(section.title, level=section.level)
        
        # Copy content while preserving formatting
        for para in section.content:
            # Skip the heading since we've already added it
            if para.text == section.title:
                continue
                
            # Skip empty paragraphs
            if not para.text.strip():
                continue
                
            new_para = new_doc.add_paragraph()
            
            # Copy runs with careful attribute handling
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                
                # Copy core run properties
                self._copy_core_run_properties(run, new_run)
                
                # Copy extended run properties
                self._copy_extended_run_properties(run, new_run)
                
                # Copy font properties
                self._copy_font_properties(run, new_run)
                
                # Copy style if it exists
                if hasattr(run, 'style') and run.style:
                    try:
                        new_run.style = run.style
                    except Exception as e:
                        print(f"Could not copy run style: {e}")
            
            # Copy paragraph style and properties
            self._copy_paragraph_properties(para, new_para)
                
        return new_doc
    
    def _copy_core_run_properties(self, source_run, target_run) -> None:
        """
        Copy the core run properties that are guaranteed to exist in python-docx.
        
        Args:
            source_run: Source run to copy from
            target_run: Target run to copy to
        """
        core_properties = ['bold', 'italic', 'underline']
        
        for prop in core_properties:
            try:
                setattr(target_run, prop, getattr(source_run, prop))
            except Exception as e:
                print(f"Could not copy core property {prop}: {e}")
    
    def _copy_extended_run_properties(self, source_run, target_run) -> None:
        """
        Copy extended run properties with validation.
        
        Args:
            source_run: Source run to copy from
            target_run: Target run to copy to
        """
        extended_properties = [
            'all_caps', 'double_strike', 'emboss', 'imprint',
            'outline', 'shadow', 'small_caps', 'strike',
            'subscript', 'superscript'
        ]
        
        for prop in extended_properties:
            try:
                if hasattr(source_run, prop):
                    setattr(target_run, prop, getattr(source_run, prop))
            except Exception as e:
                print(f"Could not copy extended property {prop}: {e}")
    
    def _copy_font_properties(self, source_run, target_run) -> None:
        """
        Copy font properties with comprehensive handling of colors and highlighting.
        
        Args:
            source_run: Source run to copy from
            target_run: Target run to copy to
        """
        if not hasattr(source_run, 'font') or not hasattr(target_run, 'font'):
            return
            
        # Copy basic font properties
        if source_run.font.name:
            target_run.font.name = source_run.font.name
            
        if source_run.font.size:
            target_run.font.size = source_run.font.size
        
        # Handle text color (foreground)
        try:
            if hasattr(source_run.font, 'color'):
                if source_run.font.color.rgb is not None:
                    target_run.font.color.rgb = source_run.font.color.rgb
                elif hasattr(source_run.font.color, 'theme_color'):
                    target_run.font.color.theme_color = source_run.font.color.theme_color
        except Exception as e:
            print(f"Could not copy font color: {e}")
            
        # Handle highlighting (background)
        try:
            if hasattr(source_run.font, 'highlight_color'):
                if source_run.font.highlight_color:
                    target_run.font.highlight_color = source_run.font.highlight_color
        except Exception as e:
            print(f"Could not copy highlight color: {e}")
            
        # Handle background color/shading using XML (for non-highlight background colors)
        try:
            # Import here to avoid dependencies at module level
            from docx.oxml.ns import qn
            
            # Get XML elements
            source_element = source_run._element
            target_element = target_run._element
            
            # Look for shading in source run's properties
            source_props = source_element.get_or_add_rPr()
            source_shd = source_props.find(qn('w:shd'))
            
            if source_shd is not None:
                # Source has custom background shading
                # Get target run properties
                target_props = target_element.get_or_add_rPr()
                
                # Check if target already has shading
                target_shd = target_props.find(qn('w:shd'))
                
                # If target already has shading, remove it (we'll copy the new one)
                if target_shd is not None:
                    target_props.remove(target_shd)
                    
                # Copy the shading properties from source to target
                from copy import deepcopy
                new_shd = deepcopy(source_shd)
                target_props.append(new_shd)
        except Exception as e:
            print(f"Could not copy background shading: {e}")
    
    def _copy_paragraph_properties(self, source_para, target_para) -> None:
        """
        Copy paragraph style, properties, and shading with validation.
        
        Args:
            source_para: Source paragraph to copy from
            target_para: Target paragraph to copy to
        """
        # Copy style if it exists
        if source_para.style:
            try:
                target_para.style = source_para.style
            except Exception as e:
                print(f"Could not copy paragraph style: {e}")
        
        # Handle paragraph shading/background
        try:
            if hasattr(source_para._element, 'pPr'):
                source_pPr = source_para._element.pPr
                if hasattr(source_pPr, 'shd'):
                    shading = source_pPr.shd
                    if shading is not None and hasattr(target_para._element, 'pPr'):
                        # Ensure pPr exists in target
                        if target_para._element.pPr is None:
                            target_para._element.get_or_add_pPr()
                        # Copy shading element
                        target_para._element.pPr.shd = shading
        except Exception as e:
            print(f"Could not copy paragraph shading: {e}")
        
        # Copy paragraph format properties if they exist
        if hasattr(source_para, 'paragraph_format') and hasattr(target_para, 'paragraph_format'):
            format_properties = [
                'alignment', 'first_line_indent', 'keep_together',
                'keep_with_next', 'left_indent', 'line_spacing',
                'right_indent', 'space_after', 'space_before'
            ]
            
            for prop in format_properties:
                try:
                    source_value = getattr(source_para.paragraph_format, prop)
                    if source_value is not None:
                        setattr(target_para.paragraph_format, prop, source_value)
                except Exception as e:
                    print(f"Could not copy paragraph format property {prop}: {e}")
    
    def _create_zip_archive(self, output_dir: Path, preserve_hierarchy: bool = False) -> Path:
        """Create zip archive with section documents, optionally preserving hierarchy."""
        zip_path = output_dir / f"{self.input_path.stem}_sections.zip"
        
        # Use temporary directory for intermediate files
        with TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as archive:
                total_sections = len(self.sections)
                for idx, section in enumerate(self.sections, 1):
                    # Check for cancellation
                    if self.cancel_requested:
                        self.status_callback("Operation canceled by user")
                        return None

                    try:
                        # Create document for section
                        doc = self._create_section_document(section)
                        
                        # Add hierarchy metadata and get metadata dict
                        metadata = self._add_document_metadata(doc, section, self.sections)
                        
                        # Create a separate metadata file
                        if metadata:
                            import json
                            # Save metadata to a JSON file with the same name + .meta.json
                            meta_file = temp_path / f"{section.safe_title}.meta.json"
                            with open(meta_file, 'w', encoding='utf-8') as f:
                                json.dump(metadata, f, indent=2)
                        
                        # Determine path within zip based on hierarchy option
                        if preserve_hierarchy and section.parent:
                            # Get folder structure from parent hierarchy
                            folder_components = section.get_path_components()
                            
                            # Create base temp file path 
                            section_dir = temp_path
                            for component in folder_components:
                                section_dir = section_dir / component
                                section_dir.mkdir(exist_ok=True, parents=True)
                            
                            temp_file = section_dir / f"{section.safe_title}.docx"
                            
                            # Create archive path with folders
                            archive_path = '/'.join(folder_components + [f"{section.safe_title}.docx"])
                        else:
                            # No hierarchy - flat structure
                            temp_file = temp_path / f"{section.safe_title}.docx"
                            archive_path = temp_file.name
                        
                        # Save to temp file
                        doc.save(temp_file)
                        
                        # Add to archive with proper path
                        archive.write(temp_file, archive_path)
                        
                        # Add metadata file to archive if it exists
                        meta_file = temp_path / f"{section.safe_title}.meta.json"
                        if meta_file.exists():
                            meta_archive_path = archive_path + ".meta.json"
                            archive.write(meta_file, meta_archive_path)
                        
                        # Report progress
                        percent_complete = int((idx / total_sections) * 100)
                        self.progress_callback(percent_complete)
                        self.status_callback(f"Processed section {idx}/{total_sections}: {section.safe_title}")
                        
                    except Exception as e:
                        self.status_callback(f"Error processing section '{section.safe_title}': {str(e)}")
                        continue
                
                if self.cancel_requested:
                    self.status_callback("Operation canceled while creating archive")
                    return None
                    
            self.status_callback(f"Created archive at: {zip_path}")
            return zip_path
    
    def _save_individual_files(self, output_dir: Path, preserve_hierarchy: bool = False) -> Path:
        """Save individual document files with optional hierarchy preservation."""
        total_sections = len(self.sections)
        files_created = 0
        
        for idx, section in enumerate(self.sections, 1):
            # Check for cancellation
            if self.cancel_requested:
                self.status_callback("Operation canceled by user")
                return None
                
            try:
                # Create document for section
                doc = self._create_section_document(section)
                
                # Add hierarchy metadata and get metadata dict
                metadata = self._add_document_metadata(doc, section, self.sections)
                
                # Create a separate metadata file
                if metadata:
                    import json
                    # Save metadata to a JSON file with the same name + .meta.json
                    meta_file_path = output_dir / f"{section.safe_title}.meta.json"
                    if preserve_hierarchy and section.parent is not None:
                        # Adjust path for hierarchical structure
                        folder_components = section.get_path_components()
                        section_dir = output_dir
                        for component in folder_components:
                            section_dir = section_dir / component
                        meta_file_path = section_dir / f"{section.safe_title}.meta.json"
                        
                    # Make sure parent directory exists
                    meta_file_path.parent.mkdir(parents=True, exist_ok=True)
                    
                    # Save metadata to file
                    with open(meta_file_path, 'w', encoding='utf-8') as f:
                        json.dump(metadata, f, indent=2)
                
                # Determine output path based on hierarchy option
                if preserve_hierarchy and section.parent is not None:
                    # Get folder path components
                    folder_components = section.get_path_components()
                    section_dir = output_dir
                    
                    # Create nested folders
                    for component in folder_components:
                        section_dir = section_dir / component
                        section_dir.mkdir(exist_ok=True, parents=True)
                    
                    output_file = section_dir / f"{section.safe_title}.docx"
                    relative_path = "/".join(folder_components + [section.safe_title])
                    print(f"Writing to hierarchical path: {relative_path}")
                else:
                    # Flat structure - just save directly to output directory
                    output_file = output_dir / f"{section.safe_title}.docx"
                    print(f"Writing to flat path: {section.safe_title}")
                
                # Save the document
                doc.save(output_file)
                files_created += 1
                
                # Report progress
                percent_complete = int((idx / total_sections) * 100)
                self.progress_callback(percent_complete)
                self.status_callback(f"Processed section {idx}/{total_sections}: {section.safe_title}")
                
            except Exception as e:
                self.status_callback(f"Error processing section '{section.safe_title}': {str(e)}")
                continue
            
        if self.cancel_requested:
            self.status_callback("Operation canceled while saving files")
            return None
                
        self.status_callback(f"Saved {files_created} documents to: {output_dir}")
        return output_dir

class DocumentSplitterDialog(QDialog):
    """
    Dialog for splitting Word documents by heading level with persistent settings.
    """
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Split Document by Headings")
        self.setMinimumSize(600, 400)
        
        # Initialize state variables
        self.input_path = None
        self.template_path = None
        self.output_dir = None
        
        # Get application settings
        self.settings = QSettings('DocxSearchApp', 'Settings')
        
        # Setup UI
        self.setup_ui()
        
        # Load previous settings
        self.load_settings()

        # Reference to the current splitter
        self.current_splitter = None
    
    def setup_ui(self):
        """Initialize the user interface."""
        layout = QVBoxLayout(self)
        
        # Input document selection
        input_group = QFrame()
        input_layout = QHBoxLayout(input_group)
        
        input_label = QLabel("Input Document:")
        self.input_field = QLineEdit()
        self.input_field.setReadOnly(True)
        self.input_field.setPlaceholderText("Select a Word document to split")
        
        input_button = QPushButton("Browse...")
        input_button.clicked.connect(self.browse_input_document)
        
        input_layout.addWidget(input_label)
        input_layout.addWidget(self.input_field, 1)
        input_layout.addWidget(input_button)
        
        layout.addWidget(input_group)
        
        # Template document selection
        template_group = QFrame()
        template_layout = QHBoxLayout(template_group)
        
        template_label = QLabel("Template Document:")
        self.template_field = QLineEdit()
        self.template_field.setReadOnly(True)
        self.template_field.setPlaceholderText("Select a template document (optional)")
        
        template_button = QPushButton("Browse...")
        template_button.clicked.connect(self.browse_template_document)
        
        template_layout.addWidget(template_label)
        template_layout.addWidget(self.template_field, 1)
        template_layout.addWidget(template_button)
        
        layout.addWidget(template_group)
        
        # Heading level selection
        level_group = QFrame()
        level_layout = QHBoxLayout(level_group)
        
        level_label = QLabel("Split at Heading Level:")
        self.level_combo = QComboBox()
        for i in range(1, 5):  # Heading levels 1-4 only
            self.level_combo.addItem(f"Heading {i}", i)
        
        # Default to heading level 3
        self.level_combo.setCurrentIndex(2)  # 0-based index, so 2 = Heading 3
        
        level_layout.addWidget(level_label)
        level_layout.addWidget(self.level_combo)
        level_layout.addStretch(1)
        
        layout.addWidget(level_group)
        
        # Output options
        output_group = QFrame()
        output_layout = QVBoxLayout(output_group)
        
        self.zip_checkbox = QCheckBox("Create ZIP archive of documents")
        self.zip_checkbox.setChecked(True)
        output_layout.addWidget(self.zip_checkbox)
        
        self.preserve_hierarchy_checkbox = QCheckBox("Preserve file organization using folders")
        self.preserve_hierarchy_checkbox.setChecked(False)
        output_layout.addWidget(self.preserve_hierarchy_checkbox)

        output_dir_frame = QFrame()
        output_dir_layout = QHBoxLayout(output_dir_frame)
        
        output_label = QLabel("Output Location:")
        self.output_field = QLineEdit()
        self.output_field.setReadOnly(True)
        self.output_field.setPlaceholderText("Select output directory")
        
        output_button = QPushButton("Browse...")
        output_button.clicked.connect(self.browse_output_directory)
        
        output_dir_layout.addWidget(output_label)
        output_dir_layout.addWidget(self.output_field, 1)
        output_dir_layout.addWidget(output_button)
        
        output_layout.addWidget(output_dir_frame)
        layout.addWidget(output_group)
        
        # Status display
        self.status_text = QLabel("Ready")
        layout.addWidget(self.status_text)
        
        # Progress bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setTextVisible(False)
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)
        
        # Buttons
        button_box = QHBoxLayout()
        
        # Process button
        self.process_button = QPushButton("Process Document")
        self.process_button.clicked.connect(self.process_document)
        button_box.addWidget(self.process_button)
        
        # Close button
        close_button = QPushButton("Close")
        close_button.clicked.connect(self.close)  # This should trigger the closeEvent
        button_box.addWidget(close_button)
        
        layout.addLayout(button_box)
    
    def load_settings(self):
        """Load previous settings."""
        # Load template path
        template_path = self.settings.value('document_splitter/template_path')
        if template_path:
            self.template_path = Path(template_path)
            self.template_field.setText(str(self.template_path))
        
        # Load output directory
        output_dir = self.settings.value('document_splitter/output_dir')
        if output_dir:
            self.output_dir = Path(output_dir)
            self.output_field.setText(str(self.output_dir))
        
        # Load heading level
        level = self.settings.value('document_splitter/heading_level', type=int)
        if level:
            # Find index of the level in the combo box
            for i in range(self.level_combo.count()):
                if self.level_combo.itemData(i) == level:
                    self.level_combo.setCurrentIndex(i)
                    break
        
        # Load zip option
        create_zip = self.settings.value('document_splitter/create_zip', type=bool)
        if create_zip is not None:  # Check for None since this is a boolean
            self.zip_checkbox.setChecked(create_zip)
    
        # Load hierarchy option
        preserve_hierarchy = self.settings.value('document_splitter/preserve_hierarchy', type=bool)
        if preserve_hierarchy is not None:
            self.preserve_hierarchy_checkbox.setChecked(preserve_hierarchy)

    def save_settings(self):
        """Save current settings."""
        # Save template path
        if self.template_path:
            self.settings.setValue('document_splitter/template_path', str(self.template_path))
        
        # Save output directory
        if self.output_dir:
            self.settings.setValue('document_splitter/output_dir', str(self.output_dir))
        
        # Save heading level
        level = self.level_combo.currentData()
        self.settings.setValue('document_splitter/heading_level', level)
        
        # Save zip option
        self.settings.setValue('document_splitter/create_zip', self.zip_checkbox.isChecked())
    
        # Save hierarchy option
        self.settings.setValue('document_splitter/preserve_hierarchy', 
                               self.preserve_hierarchy_checkbox.isChecked())

    def closeEvent(self, event):
        """Handle dialog close event."""
        print("closeEvent triggered!")  # Debug
        
        # Check if there's an active operation
        if self.current_splitter is not None and not self.process_button.isEnabled():
            print(f"Active operation detected! current_splitter={self.current_splitter}")
            reply = QMessageBox.question(
                self,
                "Cancel Operation",
                "A document split operation is in progress.\nDo you want to cancel it?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                print("User confirmed cancellation")
                # Cancel the operation
                self.update_status("Canceling operation...")
                print("About to call current_splitter.cancel()")  # Debug
                self.current_splitter.cancel()
                print("Called current_splitter.cancel()")  # Debug
                
                # Wait briefly for cancellation to take effect
                QTimer.singleShot(500, self.cleanup_and_close)
                event.ignore()  # Don't close yet, wait for cleanup
            else:
                event.ignore()  # Don't close if user doesn't want to cancel
        else:
            print(f"No active operation detected: current_splitter={self.current_splitter}, button enabled={self.process_button.isEnabled()}")
            # Save settings before closing
            self.save_settings()
            event.accept()
    
    def cleanup_and_close(self):
        """Clean up resources and close the dialog."""
        self.current_splitter = None
        self.process_button.setEnabled(True)
        self.progress_bar.setVisible(False)
        self.update_status("Operation canceled")
        self.close()  # Now we can close

    def browse_input_document(self):
        """Browse for input document."""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Document to Split",
            "",
            "Word Documents (*.docx)"
        )
        
        if file_path:
            self.input_path = Path(file_path)
            self.input_field.setText(str(self.input_path))
            
            # Set default output directory to input file's directory
            if not self.output_dir:
                self.output_dir = self.input_path.parent
                self.output_field.setText(str(self.output_dir))
                # Save this setting
                self.save_settings()
    
    def browse_template_document(self):
        """Browse for template document."""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Template Document",
            "",
            "Word Documents (*.docx)"
        )
        
        if file_path:
            self.template_path = Path(file_path)
            self.template_field.setText(str(self.template_path))
            # Save this setting
            self.save_settings()
    
    def browse_output_directory(self):
        """Browse for output directory."""
        directory = QFileDialog.getExistingDirectory(
            self,
            "Select Output Directory",
            str(self.output_dir) if self.output_dir else ""
        )
        
        if directory:
            self.output_dir = Path(directory)
            self.output_field.setText(str(self.output_dir))
            # Save this setting
            self.save_settings()
    
    def update_status(self, message):
        """Update status display."""
        self.status_text.setText(message)
        QApplication.processEvents()  # Ensure UI updates

    def update_progress(self, percent):
        """Update progress bar."""
        self.progress_bar.setValue(percent)
        QApplication.processEvents()  # Ensure UI updates

    def process_document(self):
        """Process the document with selected options."""
        # Validate inputs
        if not self.input_path or not self.input_path.exists():
            QMessageBox.warning(
                self,
                "Missing Input",
                "Please select a valid input document."
            )
            return
        
        # Save all current settings before processing
        self.save_settings()
        
        # Disable the process button during processing
        self.process_button.setEnabled(False)
        print(f"Process button disabled, enabled status={self.process_button.isEnabled()}")  # Debug
        
        if not self.template_path:
            # Create default template if none provided
            self.update_status("Creating default template document...")
            
            try:
                # Create a simple empty document as template
                default_doc = docx.Document()
                
                with TemporaryDirectory() as temp_dir:
                    temp_path = Path(temp_dir)
                    self.template_path = temp_path / "default_template.docx"
                    default_doc.save(self.template_path)
                    
                    self.split_document()
            except Exception as e:
                QMessageBox.critical(
                    self,
                    "Error",
                    f"Failed to create default template: {str(e)}"
                )
        else:
            self.split_document()
        
        # Re-enable process button after completion
        self.process_button.setEnabled(True)
    
    def split_document(self):
        """Execute the document splitting operation."""
        try:
            # Show progress bar in indeterminate mode initially
            self.progress_bar.setVisible(True)
            self.progress_bar.setRange(0, 0)  # Indeterminate mode
            
            # Get selected heading level
            level = self.level_combo.currentData()
            
            # Create splitter with both callbacks
            self.current_splitter = DocxSplitter(
                self.input_path,
                self.template_path,
                status_callback=self.update_status,
                progress_callback=self.update_progress
            )
            print(f"Current splitter created and assigned: {self.current_splitter}")  # Debug
            
            # First parse the sections - this will be in indeterminate mode
            self.update_status("Parsing document sections...")
            self.current_splitter.parse_sections(level)
            
            # Check if operation was canceled during parsing
            if self.current_splitter.cancel_requested:
                self.progress_bar.setVisible(False)
                self.process_button.setEnabled(True)
                self.current_splitter = None
                self.update_status("Operation canceled")
                return
            
            # Now we know how many sections, switch to determinate mode
            section_count = len(self.current_splitter.sections)
            if section_count > 0:
                self.progress_bar.setRange(0, 100)  # Set range to percentage
                self.progress_bar.setValue(0)       # Reset to start
                
                # Process sections with progress reporting
                create_zip = self.zip_checkbox.isChecked()
                preserve_hierarchy = self.preserve_hierarchy_checkbox.isChecked()
                
                if create_zip:
                    self.update_status(f"Creating ZIP archive with {section_count} sections...")
                else:
                    self.update_status(f"Creating {section_count} individual documents...")
                    
                result_path = self.current_splitter.process_document(
                    self.output_dir,
                    target_level=level,
                    create_zip=create_zip,
                    preserve_hierarchy=preserve_hierarchy
                )
                
                # Check if operation was canceled during processing
                if result_path is None or self.current_splitter.cancel_requested:
                    self.progress_bar.setVisible(False)
                    self.process_button.setEnabled(True)
                    self.current_splitter = None
                    self.update_status("Operation canceled")
                    return
                    
                # Ensure progress bar shows 100% at the end
                self.progress_bar.setValue(100)
                
            else:
                # No sections found
                self.update_status(f"No sections found at heading level {level}")
                result_path = self.output_dir
                
            # Hide progress bar
            self.progress_bar.setVisible(False)
            
            # Show success message (but don't close dialog)
            hierarchy_info = " with folder hierarchy preserved" if section_count > 0 and preserve_hierarchy else ""
            msg_text = (
                f"Document processed.\n"
                f"Found {section_count} sections at heading level {level}{hierarchy_info}.\n\n"
                f"Output saved to: {result_path}"
            )
            
            QMessageBox.information(
                self,
                "Success",
                msg_text
            )
            
            # Clear input field to encourage selecting a new document
            self.input_path = None
            self.input_field.clear()
            self.update_status("Ready for next document")
            
            # Clear current_splitter reference at the end
            self.current_splitter = None
            self.process_button.setEnabled(True)
            
        except Exception as e:
            # Hide progress bar
            self.progress_bar.setVisible(False)
            self.current_splitter = None
            self.process_button.setEnabled(True)
            
            QMessageBox.critical(
                self,
                "Error",
                f"Failed to split document: {str(e)}"
            )

class PrefixManager:
    """
    Manages search prefixes and their associated folder mappings.
    Handles validation, storage, and retrieval of prefix configurations.
    """
    
    def __init__(self, settings: QSettings):
        self.settings = settings
        self.prefix_configs: Dict[str, Set[str]] = {}
        self.excluded_folders: Set[str] = set()
        self._load_prefixes()
    
    def _is_valid_prefix(self, prefix: str) -> bool:
        """
        Validate a prefix string.
        Only allows alphanumeric characters (no spaces or special characters).
        """
        return bool(re.match(r'^[a-zA-Z0-9]+$', prefix))
    
    def _load_prefixes(self) -> None:
        """Load prefix configurations from settings."""
        self.prefix_configs.clear()
        size = self.settings.beginReadArray("prefixes")
        
        for i in range(size):
            self.settings.setArrayIndex(i)
            prefix = self.settings.value("prefix", "")
            folders = self.settings.value("folders", set())
            
            # Convert folders to set if it's not already
            if isinstance(folders, str):
                folders = {folders}
            elif isinstance(folders, list):
                folders = set(folders)
                
            if self._is_valid_prefix(prefix):
                self.prefix_configs[prefix] = folders
                
        self.settings.endArray()

        excluded_folders = self.settings.value("excluded_folders", [])
        if excluded_folders:
            if isinstance(excluded_folders, str):
                self.excluded_folders = {excluded_folders}
            else:
                self.excluded_folders = set(excluded_folders)

    def _save_prefixes(self) -> None:
        """Save prefix configurations to settings."""
        self.settings.beginWriteArray("prefixes")
        
        for i, (prefix, folders) in enumerate(self.prefix_configs.items()):
            self.settings.setArrayIndex(i)
            self.settings.setValue("prefix", prefix)
            self.settings.setValue("folders", list(folders))
            
        self.settings.endArray()

        self.settings.setValue("excluded_folders", list(self.excluded_folders))
    
    def add_prefix_folder(self, prefix: str, folder: str) -> bool:
        """
        Add a folder to a prefix's configuration.
        Creates new prefix if it doesn't exist.
        
        Args:
            prefix: The prefix to configure
            folder: Path to the folder to associate with the prefix
            
        Returns:
            bool: True if successful, False if prefix is invalid
        """
        if not self._is_valid_prefix(prefix):
            return False
            
        if prefix not in self.prefix_configs:
            self.prefix_configs[prefix] = set()
            
        self.prefix_configs[prefix].add(folder)
        self._save_prefixes()
        return True
    
    def remove_prefix_folder(self, prefix: str, folder: str) -> bool:
        """
        Remove a folder from a prefix's configuration.
        
        Args:
            prefix: The prefix to modify
            folder: Path to remove from the prefix
            
        Returns:
            bool: True if successful, False if prefix doesn't exist
        """
        if prefix not in self.prefix_configs:
            return False
            
        self.prefix_configs[prefix].discard(folder)
        
        # Remove prefix entirely if no folders remain
        if not self.prefix_configs[prefix]:
            del self.prefix_configs[prefix]
            
        self._save_prefixes()
        return True
    
    def get_folders_for_prefix(self, prefix: str) -> Set[str]:
        """Get all folders associated with a prefix."""
        return self.prefix_configs.get(prefix, set()).copy()
    
    def is_valid_prefix_word(self, word: str) -> bool:
        """Check if a word is a configured prefix."""
        return word in self.prefix_configs
    
    def export_to_csv(self, filepath: str) -> bool:
        """
        Export prefix configurations to CSV file.
        
        Format:
        prefix,folder1|folder2|folder3
        """
        try:
            with open(filepath, 'w', newline='') as f:
                writer = csv.writer(f)
                writer.writerow(['prefix', 'folders'])  # Header
                
                for prefix, folders in self.prefix_configs.items():
                    writer.writerow([prefix, '|'.join(folders)])
            return True
        except Exception as e:
            print(f"Error exporting prefixes: {e}")
            return False
    
    def import_from_csv(self, filepath: str) -> bool:
        """
        Import prefix configurations from CSV file.
        
        Expected format:
        prefix,folder1|folder2|folder3
        """
        try:
            new_configs = {}
            with open(filepath, 'r', newline='') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    prefix = row['prefix']
                    if self._is_valid_prefix(prefix):
                        folders = set(row['folders'].split('|'))
                        new_configs[prefix] = folders
            
            self.prefix_configs = new_configs
            self._save_prefixes()
            return True
        except Exception as e:
            print(f"Error importing prefixes: {e}")
            return False
    
    def check_folder_exists(self, base_path: str, folder: str) -> bool:
        """
        Check if a configured folder still exists.
        
        Args:
            base_path: Base search path
            folder: Relative folder path to check
            
        Returns:
            bool: True if folder exists
        """
        full_path = Path(base_path) / folder
        return full_path.exists() and full_path.is_dir()
    
    def verify_folders_exist(self, base_path: str) -> List[tuple[str, str]]:
        """
        Verify all configured folders exist.
        
        Args:
            base_path: Base search path
            
        Returns:
            List of (prefix, folder) pairs for missing folders
        """
        missing = []
        for prefix, folders in self.prefix_configs.items():
            for folder in folders:
                if not self.check_folder_exists(base_path, folder):
                    missing.append((prefix, folder))
        return missing

    def is_folder_excluded(self, folder: str) -> bool:
        """
        Check if a folder is excluded from general searches.
        """
        # Check if folder itself is excluded (direct match)
        if folder in self.excluded_folders:
            return True
        
        # Check if folder is a subfolder of any excluded folder
        for excluded in self.excluded_folders:
            # Skip empty folders
            if not excluded:
                continue
            # Check if this is a subfolder (using path separator handling)
            if folder.startswith(excluded + '/') or folder.startswith(excluded + '\\'):
                return True
        
        # If we reach here, the folder is not excluded
        return False
    
    def set_folder_exclusion(self, folder: str, excluded: bool) -> None:
        """
        Set whether a folder should be excluded from general searches.
        Handles parent-child folder relationships intelligently.
        """
        if excluded:
            # Check if a parent folder is already excluded
            for excluded_folder in self.excluded_folders:
                if not excluded_folder:
                    continue
                if folder.startswith(excluded_folder + '/') or folder.startswith(excluded_folder + '\\'):
                    # Parent is already excluded, this folder is implicitly excluded
                    return
            
            # Add this folder to exclusions
            self.excluded_folders.add(folder)
            
            # Remove any subfolders from exclusions as they're now implicitly excluded
            subfolders_to_remove = []
            for excluded_folder in self.excluded_folders:
                if not excluded_folder or excluded_folder == folder:
                    continue
                if excluded_folder.startswith(folder + '/') or excluded_folder.startswith(folder + '\\'):
                    subfolders_to_remove.append(excluded_folder)
            
            for subfolder in subfolders_to_remove:
                self.excluded_folders.discard(subfolder)
        else:
            # Just remove the folder from exclusions
            self.excluded_folders.discard(folder)
            
        self._save_prefixes()

class ShortcutDialog(QDialog):
    """
    Dialog for configuring global keyboard shortcuts.
    
    Uses a QLineEdit in read-only mode to capture keystrokes, converting them
    into a format suitable for the keyboard library while providing a clean UI
    for shortcut configuration.
    """
    
    def __init__(self, current_shortcut: str, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Set Window Focus Shortcut")
        self.setModal(True)
        
        # Store the shortcut in keyboard library format
        self.shortcut = current_shortcut
        
        # Setup UI
        layout = QVBoxLayout(self)
        
        # Instructions
        instructions = QLabel(
            "Press the desired key combination.\n"
            "The shortcut will be updated when you click OK."
        )
        instructions.setWordWrap(True)
        layout.addWidget(instructions)
        
        # Shortcut input field
        self.shortcut_input = QLineEdit()
        self.shortcut_input.setReadOnly(True)
        self.shortcut_input.setText(self._format_shortcut_for_display(current_shortcut))
        layout.addWidget(self.shortcut_input)
        
        # Dialog buttons
        button_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | 
            QDialogButtonBox.StandardButton.Cancel
        )
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)
        
        # Initialize with empty modifiers
        self.current_modifiers = set()
    
    def _format_shortcut_for_display(self, shortcut: str) -> str:
        """Convert keyboard library format to display format."""
        # Split combination into parts
        parts = shortcut.split('+')
        # Capitalize each part for display
        return ' + '.join(part.capitalize() for part in parts)
    
    def _format_shortcut_for_storage(self, modifiers: set, key: str) -> str:
        """Convert captured keystroke into keyboard library format."""
        parts = list(modifiers) + [key.lower()]
        return '+'.join(parts)
    
    def keyPressEvent(self, event: QKeyEvent) -> None:
        """
        Handle keystroke capture with comprehensive key support.
        
        Maps Qt key codes to keyboard library compatible strings while preventing
        event propagation. Supports standard keys, special characters, and modifiers
        with precise Qt constant mappings.
        """
        # Prevent event propagation
        event.accept()
        
        # Only process new keystrokes, not auto-repeats
        if not event.isAutoRepeat():
            self.current_modifiers.clear()
            
            # Map modifiers with precise Qt constants
            if event.modifiers() & Qt.KeyboardModifier.ControlModifier:
                self.current_modifiers.add('ctrl')
            if event.modifiers() & Qt.KeyboardModifier.AltModifier:
                self.current_modifiers.add('alt')
            if event.modifiers() & Qt.KeyboardModifier.ShiftModifier:
                self.current_modifiers.add('shift')
            
            # Get the main key
            key = None
            
            # Handle different key categories with precise Qt constants
            if Qt.Key.Key_A <= event.key() <= Qt.Key.Key_Z:
                # Letter keys: map to lowercase
                key = chr(event.key()).lower()
            elif Qt.Key.Key_0 <= event.key() <= Qt.Key.Key_9:
                # Number keys: direct mapping
                key = chr(event.key())
            elif Qt.Key.Key_F1 <= event.key() <= Qt.Key.Key_F12:
                # Function keys: calculate F-key number
                key = f'f{event.key() - Qt.Key.Key_F1 + 1}'
            else:
                # Special keys mapping using verified Qt constants
                key = {
                    Qt.Key.Key_Space: 'space',
                    Qt.Key.Key_Tab: 'tab',
                    Qt.Key.Key_Return: 'enter',
                    Qt.Key.Key_Enter: 'enter',
                    Qt.Key.Key_Backspace: 'backspace',
                    Qt.Key.Key_Delete: 'delete',
                    Qt.Key.Key_Escape: 'esc',
                    Qt.Key.Key_Up: 'up',
                    Qt.Key.Key_Down: 'down',
                    Qt.Key.Key_Left: 'left',
                    Qt.Key.Key_Right: 'right',
                    Qt.Key.Key_QuoteLeft: '`',  # Backtick key
                    Qt.Key.Key_Minus: '-',
                    Qt.Key.Key_Equal: '=',
                    Qt.Key.Key_BracketLeft: '[',
                    Qt.Key.Key_BracketRight: ']',
                    Qt.Key.Key_Semicolon: ';',
                    Qt.Key.Key_Apostrophe: "'",  # Corrected from Key_Quote
                    Qt.Key.Key_Comma: ',',
                    Qt.Key.Key_Period: '.',
                    Qt.Key.Key_Slash: '/',
                    Qt.Key.Key_Backslash: '\\',
                }.get(event.key())
            
            # Update shortcut if we have a valid key combination
            if key and (self.current_modifiers or key in {f'f{i}' for i in range(1, 13)}):
                self.shortcut = self._format_shortcut_for_storage(
                    self.current_modifiers, key
                )
                self.shortcut_input.setText(
                    self._format_shortcut_for_display(self.shortcut)
                )

class SearchResultItem(QListWidgetItem):
    """
    Enhanced list item for displaying search results with folder context.
    
    This widget provides:
    1. Sophisticated document metadata display
    2. Folder path context when relevant
    3. Visual hierarchy between filename and folder path
    """
    
    def __init__(self, doc_info: DocumentInfo):
        # Initialize with empty text - we'll set it in format_display
        super().__init__()
        self.doc_info = doc_info
        self.format_display()
        
        # Store the full path as item data for activation handling
        self.setData(Qt.ItemDataRole.UserRole, str(doc_info.path))
    
    def format_display(self):
        """
        Format the display text with intelligent folder path handling.
        
        Creates a visually hierarchical display with:
        - Filename as primary information
        - Folder path as secondary context when available
        """
        display_text = self.doc_info.name
        
        # Add folder context if available
        if self.doc_info.relative_path:
            display_text = f"{display_text}  [{self.doc_info.relative_path}]"
        
        self.setText(display_text)
        
        # Optional: Add tooltip with full metadata
        tooltip = (
            f"Name: {self.doc_info.name}\n"
            f"Path: {self.doc_info.relative_path}\n"
            f"Size: {self.doc_info.size:,} bytes\n"
            f"Modified: {datetime.fromtimestamp(self.doc_info.last_modified).strftime('%Y-%m-%d %H:%M:%S')}"
        )
        
        # Add hierarchy info to tooltip if available
        if hasattr(self.doc_info, 'original_doc_path') and self.doc_info.original_doc_path:
            tooltip += f"\nFrom document: {Path(self.doc_info.original_doc_path).name}"
            
        self.setToolTip(tooltip)

class SearchInputWithKeyNavigation(QLineEdit):
    """
    Enhanced QLineEdit that intelligently handles keyboard navigation and prefix highlighting.
    """
    
    ctrlEnterPressed = pyqtSignal()
    
    def __init__(self, results_list: QListWidget, prefix_manager: PrefixManager, parent=None):
        super().__init__(parent)
        self.results_list = results_list
        self.prefix_manager = prefix_manager
        
        # Cache for efficient prefix validation
        self.current_prefix = None
        self.prefix_valid = False
        
        # Style configurations
        self.default_style = self.styleSheet()
        self.prefix_style = """
            QLineEdit {
                background-color: #2e7d32;
                color: white;
                padding: 2px 5px;
                border: 1px solid #1b5e20;
            }
        """
        
        # Keys that should be forwarded to the results list
        self.navigation_keys = {
            Qt.Key.Key_Up,
            Qt.Key.Key_Down,
            Qt.Key.Key_PageUp,
            Qt.Key.Key_PageDown
        }
        
        # Connect text change handler
        self.textChanged.connect(self._handle_text_changed)

    def _handle_text_changed(self, text: str):
        """Handle text changes with prefix detection and styling."""
        self.current_prefix = None
        self.prefix_valid = False
        
        # Reset styling if empty
        if not text:
            self.setStyleSheet(self.default_style)
            return
            
        # Check for potential prefix (word followed by space)
        parts = text.strip().split(maxsplit=1)
        if len(parts) < 2:
            self.setStyleSheet(self.default_style)
            return
            
        potential_prefix = parts[0]
        if self.prefix_manager.is_valid_prefix_word(potential_prefix):
            self.current_prefix = potential_prefix
            self.prefix_valid = True
            self.setStyleSheet(self.prefix_style)
        else:
            self.setStyleSheet(self.default_style)

    def keyPressEvent(self, event: QKeyEvent) -> None:
        """
        Handle key press events with special handling for Ctrl+Enter.
        """
        # Print debug info for key presses - can be removed later
        print(f"Key pressed: {event.key()}, Modifiers: {event.modifiers()}")
        
        # Handle Ctrl+Enter specifically
        if ((event.key() == Qt.Key.Key_Return or event.key() == Qt.Key.Key_Enter) and 
            event.modifiers() == Qt.KeyboardModifier.ControlModifier):
            print("Ctrl+Enter detected")
            self.ctrlEnterPressed.emit()
            event.accept()
            return
            
        if event.key() in self.navigation_keys:
            # Forward the event to the results list and transfer focus
            self.results_list.setFocus()
            
            # For Down key, just focus the current selection in the results list
            if event.key() == Qt.Key.Key_Down and self.results_list.count() > 0:
                # If there's no current item, select the first one
                if self.results_list.currentRow() == -1:
                    self.results_list.setCurrentRow(0)
                
                # Make sure the item is visible
                if self.results_list.currentItem():
                    self.results_list.scrollToItem(
                        self.results_list.currentItem(),
                        QListWidget.ScrollHint.EnsureVisible
                    )
                
                # Event has been handled
                event.accept()
                return
            
            # For other navigation keys, forward the event
            new_event = QKeyEvent(
                QEvent.Type.KeyPress,
                event.key(),
                event.modifiers(),
                event.text(),
                event.isAutoRepeat(),
                event.count()
            )
            self.results_list.keyPressEvent(new_event)
            # Do not return focus to search input - we want to keep focus on the results list
        else:
            # Handle all other keys normally
            super().keyPressEvent(event)
    
    def get_current_prefix(self) -> Optional[str]:
        """Get the currently active prefix if any is valid."""
        return self.current_prefix if self.prefix_valid else None

class PrefixManagerDialog(QDialog):
    """
    Dialog for managing prefix-to-folder mappings with sophisticated validation.
    
    This dialog provides:
    1. Interactive table view of prefix configurations
    2. Real-time validation of prefix inputs
    3. Folder selection integration
    4. Multi-folder support per prefix
    """
    
    def __init__(self, prefix_manager: PrefixManager, base_path: str, parent=None):
        super().__init__(parent)
        self.prefix_manager = prefix_manager
        self.base_path = Path(base_path)
        
        self.setWindowTitle("Prefix Configuration")
        self.setModal(True)  # Modal dialog for focused interaction
        self.setMinimumSize(800, 400)  # Wider dialog for better button layout
        
        self.setup_ui()
        self.load_current_mappings()

    def setup_ui(self):
        """Initialize the user interface with sophisticated layout management."""
        layout = QVBoxLayout(self)
        
        # Instructions label
        instructions = QLabel(
            "Configure prefixes to limit searches to specific folders.\n"
            "Prefixes must be alphanumeric (no spaces or special characters)."
        )
        instructions.setWordWrap(True)
        layout.addWidget(instructions)
        
        # Table for displaying mappings
        self.mappings_table = QTableWidget()
        self.mappings_table.setColumnCount(4)  # Now 4 columns instead of 3
        self.mappings_table.setHorizontalHeaderLabels(['Prefix', 'Folders', 'Exclude from General', 'Actions'])
        
        # Set column widths
        self.mappings_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Fixed)
        self.mappings_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        self.mappings_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.Fixed)
        self.mappings_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.Fixed)
        self.mappings_table.setColumnWidth(0, 100)  # Prefix column width
        self.mappings_table.setColumnWidth(2, 150)  # Exclude column width
        self.mappings_table.setColumnWidth(3, 300)  # Actions column width
        
        self.mappings_table.setSelectionMode(QTableWidget.SelectionMode.NoSelection)
        layout.addWidget(self.mappings_table)
        
        # Add New Mapping section
        add_frame = QFrame()
        add_layout = QHBoxLayout(add_frame)
        
        self.new_prefix_input = QLineEdit()
        self.new_prefix_input.setPlaceholderText("Enter new prefix")
        self.new_prefix_input.setMaxLength(20)  # Reasonable limit for prefix length
        add_layout.addWidget(self.new_prefix_input)
        
        select_folder_btn = QPushButton("Select Folder")
        select_folder_btn.setMinimumWidth(90)
        select_folder_btn.clicked.connect(self.select_folder_for_new_prefix)
        add_layout.addWidget(select_folder_btn)
        
        layout.addWidget(add_frame)
        
        # Buttons
        button_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | 
            QDialogButtonBox.StandardButton.Cancel
        )
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)

    def load_current_mappings(self):
        """Load and display current prefix configurations."""
        self.mappings_table.setRowCount(0)  # Clear existing rows
        
        for prefix, folders in self.prefix_manager.prefix_configs.items():
            self.add_mapping_row(prefix, folders)
    
    def add_mapping_row(self, prefix: str, folders: Set[str]):
        """
        Add a new row to the mappings table with sophisticated widget management.
        
        Args:
            prefix: The prefix string
            folders: Set of folder paths associated with the prefix
        """
        row = self.mappings_table.rowCount()
        self.mappings_table.insertRow(row)
        
        # Prefix cell
        prefix_item = QTableWidgetItem(prefix)
        prefix_item.setFlags(prefix_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
        self.mappings_table.setItem(row, 0, prefix_item)
        
        # Folders cell - display as comma-separated list
        folders_text = ", ".join(sorted(folders))
        folders_item = QTableWidgetItem(folders_text)
        folders_item.setFlags(folders_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
        self.mappings_table.setItem(row, 1, folders_item)
        
        # Exclude cell - creates checkboxes for each folder
        exclude_widget = QWidget()
        exclude_layout = QVBoxLayout(exclude_widget)
        exclude_layout.setContentsMargins(4, 4, 4, 4)
        exclude_layout.setSpacing(4)
        
        for folder in sorted(folders):
            checkbox = QCheckBox()
            checkbox.setChecked(self.prefix_manager.is_folder_excluded(folder))
            checkbox.setToolTip(f"Exclude {folder} and all its subfolders from general searches")
            # Use folder name as tooltip instead of checkbox text for cleaner display
            checkbox.setToolTip(folder)
            checkbox.stateChanged.connect(lambda state, f=folder: self.toggle_folder_exclusion(f, state))
            
            # Create a folder label for better display
            folder_label = QLabel(folder.split('/')[-1])
            folder_label.setToolTip(folder)
            
            # Add to a horizontal layout
            folder_row = QHBoxLayout()
            folder_row.addWidget(checkbox)
            folder_row.addWidget(folder_label)
            folder_row.addStretch()
            
            exclude_layout.addLayout(folder_row)
        
        exclude_layout.addStretch()
        self.mappings_table.setCellWidget(row, 2, exclude_widget)
        
        # Actions cell
        actions_widget = QWidget()
        actions_layout = QHBoxLayout(actions_widget)
        actions_layout.setContentsMargins(4, 4, 4, 4)
        actions_layout.setSpacing(6)  # Add spacing between buttons
        
        add_folder_btn = QPushButton("Add Folder")
        add_folder_btn.setMinimumWidth(90)
        add_folder_btn.clicked.connect(lambda: self.add_folder_to_prefix(prefix))
        actions_layout.addWidget(add_folder_btn)
        
        remove_folder_btn = QPushButton("Remove Folder")
        remove_folder_btn.setMinimumWidth(90)
        remove_folder_btn.clicked.connect(lambda: self.remove_folder_from_prefix(prefix))
        actions_layout.addWidget(remove_folder_btn)
        
        delete_btn = QPushButton("Delete")
        delete_btn.setMinimumWidth(90)
        delete_btn.clicked.connect(lambda: self.delete_prefix(prefix))
        actions_layout.addWidget(delete_btn)
        
        self.mappings_table.setCellWidget(row, 3, actions_widget)
        
        # Adjust row height to accommodate buttons and checkboxes
        self.mappings_table.setRowHeight(row, max(actions_widget.sizeHint().height() + 8, 
                                                 exclude_widget.sizeHint().height() + 8))

    def toggle_folder_exclusion(self, folder: str, state: int):
        """Toggle exclusion state for a folder."""
        # Convert Qt.CheckState to boolean (2 = Qt.Checked, 0 = Qt.Unchecked)
        excluded = state == Qt.CheckState.Checked.value
        self.prefix_manager.set_folder_exclusion(folder, excluded)

    def select_folder_for_new_prefix(self):
        """Handle folder selection for new prefix with validation."""
        prefix = self.new_prefix_input.text().strip()
        
        if not prefix:
            QMessageBox.warning(
                self,
                "Invalid Prefix",
                "Please enter a prefix before selecting a folder."
            )
            return
        
        if not self.prefix_manager._is_valid_prefix(prefix):
            QMessageBox.warning(
                self,
                "Invalid Prefix",
                "Prefix must contain only alphanumeric characters (no spaces)."
            )
            return
        
        # Show folder selection dialog
        folder = QFileDialog.getExistingDirectory(
            self,
            "Select Folder for Prefix",
            str(self.base_path),
            QFileDialog.Option.ShowDirsOnly
        )
        
        if folder:
            try:
                rel_path = str(Path(folder).relative_to(self.base_path))
                
                # Add new prefix-folder mapping
                self.prefix_manager.add_prefix_folder(prefix, rel_path)
                
                # Explicitly ensure it's not excluded by default
                self.prefix_manager.set_folder_exclusion(rel_path, False)
                
                # Clear input and refresh display
                self.new_prefix_input.clear()
                self.load_current_mappings()
                
            except ValueError:
                QMessageBox.warning(
                    self,
                    "Invalid Folder",
                    "Selected folder must be within the base search directory."
                )

    def add_folder_to_prefix(self, prefix: str):
        """Add another folder to an existing prefix."""
        folder = QFileDialog.getExistingDirectory(
            self,
            f"Select Additional Folder for Prefix '{prefix}'",
            str(self.base_path),
            QFileDialog.Option.ShowDirsOnly
        )
        
        if folder:
            try:
                rel_path = str(Path(folder).relative_to(self.base_path))
                
                # Add new prefix-folder mapping
                self.prefix_manager.add_prefix_folder(prefix, rel_path)
                
                # Explicitly ensure it's not excluded by default
                self.prefix_manager.set_folder_exclusion(rel_path, False)
                
                # Clear input and refresh display
                self.load_current_mappings()
            except ValueError:
                QMessageBox.warning(
                    self,
                    "Invalid Folder",
                    "Selected folder must be within the base search directory."
                )

    def remove_folder_from_prefix(self, prefix: str):
        """Remove a folder from a prefix with user selection."""
        folders = sorted(self.prefix_manager.get_folders_for_prefix(prefix))
        if not folders:
            return
        
        # If only one folder, remove it directly
        if len(folders) == 1:
            self.prefix_manager.remove_prefix_folder(prefix, folders[0])
            self.load_current_mappings()
            return
        
        # For multiple folders, show selection dialog
        folder, ok = QInputDialog.getItem(
            self,
            f"Remove Folder from Prefix '{prefix}'",
            "Select folder to remove:",
            folders,
            0,
            False
        )
        
        if ok and folder:
            self.prefix_manager.remove_prefix_folder(prefix, folder)
            self.load_current_mappings()

    def delete_prefix(self, prefix: str):
        """Delete a prefix and all its folder mappings."""
        reply = QMessageBox.question(
            self,
            "Confirm Deletion",
            f"Are you sure you want to delete the prefix '{prefix}' and all its folder mappings?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            # Get all folders and remove them
            folders = self.prefix_manager.get_folders_for_prefix(prefix)
            for folder in folders:
                self.prefix_manager.remove_prefix_folder(prefix, folder)
            
            self.load_current_mappings()

    def accept(self):
        """Handle dialog acceptance with validation."""
        # Verify all folders still exist
        missing = self.prefix_manager.verify_folders_exist(self.base_path)
        if missing:
            message = "The following folders no longer exist:\n\n"
            for prefix, folder in missing:
                message += f"Prefix '{prefix}' -> folder '{folder}'\n"
            message += "\nDo you want to remove these invalid mappings?"
            
            reply = QMessageBox.question(
                self,
                "Missing Folders",
                message,
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                for prefix, folder in missing:
                    self.prefix_manager.remove_prefix_folder(prefix, folder)
        
        super().accept()

class DocumentContextList(QListWidget):
    """
    List widget for displaying document context (related documents from the same source).
    """
    itemActivated = pyqtSignal(QListWidgetItem)  # Standard activation (Enter key)
    ctrlEnterPressed = pyqtSignal(QListWidgetItem)  # For alternate paste mode
    closeContextView = pyqtSignal()  # Signal to close the context view (Right arrow)
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFocusPolicy(Qt.FocusPolicy.StrongFocus)
        
    def keyPressEvent(self, event: QKeyEvent) -> None:
        """Handle key press events with special handling for navigation."""
        # Handle standard activation with Enter
        if (event.key() == Qt.Key.Key_Return or event.key() == Qt.Key.Key_Enter) and event.modifiers() == Qt.KeyboardModifier.NoModifier:
            current_item = self.currentItem()
            if current_item:
                self.itemActivated.emit(current_item)
                event.accept()
                return
                
        # Handle Ctrl+Enter for alternate paste mode
        elif ((event.key() == Qt.Key.Key_Return or event.key() == Qt.Key.Key_Enter) and 
              event.modifiers() == Qt.KeyboardModifier.ControlModifier):
            current_item = self.currentItem()
            if current_item:
                self.ctrlEnterPressed.emit(current_item)
                event.accept()
                return
        
        # Handle Right arrow to exit context view
        elif event.key() == Qt.Key.Key_Right and event.modifiers() == Qt.KeyboardModifier.NoModifier:
            self.closeContextView.emit()
            event.accept()
            return
            
        # Handle all other keys normally
        super().keyPressEvent(event)

class EnhancedResultsList(QListWidget):
    """
    Enhanced list widget that handles Ctrl+Enter for alternate paste mode
    and Left arrow for showing document context.
    """
    ctrlEnterPressed = pyqtSignal(QListWidgetItem)
    showDocumentContext = pyqtSignal(DocumentInfo)  # New signal for context view
    focusSearchInput = pyqtSignal()  # Signal to request focus back to search input
    
    def keyPressEvent(self, event: QKeyEvent) -> None:
        """Handle key press events with special handling for navigation."""
        # Handle Ctrl+Enter for alternate paste mode
        if ((event.key() == Qt.Key.Key_Return or event.key() == Qt.Key.Key_Enter) and 
            event.modifiers() == Qt.KeyboardModifier.ControlModifier):
            current_item = self.currentItem()
            if current_item:
                self.ctrlEnterPressed.emit(current_item)
                event.accept()
                return
        
        # Handle Left arrow for showing document context
        elif event.key() == Qt.Key.Key_Left and event.modifiers() == Qt.KeyboardModifier.NoModifier:
            current_item = self.currentItem()
            if current_item and hasattr(current_item, 'doc_info'):
                self.showDocumentContext.emit(current_item.doc_info)
                event.accept()
                return
        
        # Handle Up key when at the first item - return focus to search input
        elif event.key() == Qt.Key.Key_Up and self.currentRow() == 0:
            self.focusSearchInput.emit()  # Request focus back to search input
            event.accept()
            return
                
        # Handle all other keys normally
        super().keyPressEvent(event)

class WordHandler:
    """
    Manages Word automation with sophisticated state management and error handling.
    
    This class provides a robust interface for Word operations, carefully managing
    COM object lifecycles and application state to prevent interference with
    user's existing Word sessions. It supports multiple paste modes and implements
    careful state preservation during document operations.
    """
    
    def __init__(self):
        self._app_state: Dict[str, Any] = {}
    
    @contextmanager
    def word_session(self):
        """
        Context manager for Word application sessions with state preservation.
        
        Carefully manages COM initialization and cleanup while preserving
        existing Word application state. Uses a sophisticated state tracking
        system to ensure reliable cleanup even in error cases.
        """
        pythoncom.CoInitialize()
        try:
            try:
                word_app = win32com.client.GetActiveObject("Word.Application")
            except pythoncom.com_error:
                word_app = win32com.client.Dispatch("Word.Application")
            
            # Store and optimize display settings
            self._store_application_state(word_app)
            self._optimize_application_display(word_app)
            
            yield word_app
            
        finally:
            if 'word_app' in locals():
                self._restore_application_state(word_app)
                try:
                    word_app.ScreenUpdating = True
                except:
                    pass

            pythoncom.CoUninitialize()
    
    def _store_application_state(self, word_app):
        """Store essential application display settings."""
        self._app_state = {
            'screen_updating': word_app.ScreenUpdating,
            'display_alerts': word_app.DisplayAlerts
        }
    
    def _optimize_application_display(self, word_app):
        """Optimize display settings for background operation."""
        word_app.ScreenUpdating = False
        word_app.DisplayAlerts = False
    
    def _restore_application_state(self, word_app):
        """Carefully restore previous application state."""
        for prop, value in self._app_state.items():
            try:
                if value is not None:
                    setattr(word_app, prop, value)
            except Exception:
                continue

    @contextmanager
    def open_document(self, word_app, file_path: str, readonly: bool = True):
        """
        Context manager for document operations with proper state management.
        
        Args:
            word_app: Word application instance
            file_path: Path to the document
            readonly: Whether to open in readonly mode (default: True)
        """
        doc = None
        try:
            abs_path = str(Path(file_path).resolve())
            doc = word_app.Documents.Open(
                abs_path,
                ReadOnly=readonly,
                AddToRecentFiles=False,
                NoEncodingDialog=True
            )
            yield doc
        finally:
            if doc is not None:
                try:
                    doc.Close(SaveChanges=not readonly)
                except Exception:
                    pass
    
    def get_active_documents(self) -> List[ActiveDocument]:
        """
        Enumerate currently open documents in Word with robust COM property access.
        
        Returns a list of ActiveDocument objects representing all open documents,
        carefully handling COM interaction to prevent state corruption.
        """
        active_docs = []
        try:
            with self.word_session() as word_app:
                # First verify we have documents open
                if not word_app.Documents.Count:
                    return []

                # Enumerate all open documents with careful property access
                for i, doc in enumerate(word_app.Documents, start=1):
                    try:
                        # Get document name (most reliable property)
                        name = doc.Name

                        # Get full path with error handling
                        try:
                            path = doc.FullName
                        except Exception:
                            path = name  # Fallback to name if path unavailable

                        # Generate a unique identifier using multiple approaches
                        doc_id = None
                        try:
                            # Try getting built-in UniqueID first
                            doc_id = str(doc.UniqueID)
                        except Exception:
                            try:
                                # Fallback to hash of document properties
                                doc_id = f"{name}_{path}_{doc.Saved}_{doc.Windows.Count}"
                            except Exception:
                                # Last resort: use combination of index and name
                                doc_id = f"doc_{i}_{name}"

                        active_docs.append(ActiveDocument(
                            name=name,
                            path=path,
                            window_index=i,
                            doc_id=doc_id
                        ))
                    except Exception as e:
                        print(f"Error accessing document {i}: {str(e)}")
                        continue
                try:
                    word_app.ScreenUpdating = True
                except:
                    pass

        except Exception as e:
            print(f"Error enumerating documents: {str(e)}")
            try:
                # Create a new connection just to restore screen updating
                word_app = win32com.client.GetActiveObject("Word.Application")
                word_app.ScreenUpdating = True
            except:
                pass
        
        return active_docs

    def copy_to_clipboard(self, file_path: str) -> bool:
        """Copy document content to clipboard."""
        try:
            with self.word_session() as word_app:
                with self.open_document(word_app, file_path) as doc:
                    doc.Content.Copy()
                return True
        except Exception as e:
            print(f"Error copying to clipboard: {e}")
            return False

    def transfer_content(self, source_path: str, target_path: str) -> bool:
        """
        Transfer content between documents with formatting preservation.
        
        Args:
            source_path: Path to source document
            target_path: Path to target document
            
        Returns:
            bool: True if transfer succeeded, False otherwise
        """
        try:
            with self.word_session() as word_app:
                # Copy content from source
                with self.open_document(word_app, source_path) as source_doc:
                    source_doc.Content.Copy()
                
                # Paste to target with formatting
                with self.open_document(word_app, target_path, readonly=False) as target_doc:
                    end_point = target_doc.Content.End - 1
                    target_doc.Range(end_point, end_point).Select()
                    
                    # Add spacing for content separation
                    selection = word_app.Selection
                    selection.InsertParagraphBefore()
                    selection.InsertParagraphBefore()
                    selection.Collapse(0)  # Collapse to end
                    
                    # Paste with original formatting
                    selection.PasteAndFormat(16)  # wdFormatOriginalFormatting
                
                return True
        
        except Exception as e:
            print(f"Error transferring content: {e}")
            return False

    def paste_to_active_document(self, source_path: str, target_doc_id: str, mode: str = PasteMode.CURSOR) -> bool:
        """
        Paste content at current cursor position in target document, with precise cursor positioning.
        
        Args:
            source_path: Path to source document
            target_doc_id: Composite identifier of target document
            mode: Paste mode (cursor or end position)
            
        Returns:
            bool: True if paste operation succeeded
        """
        try:
            with self.word_session() as word_app:
                # First, copy source content
                with self.open_document(word_app, source_path) as source_doc:
                    source_doc.Content.Copy()
                
                # Find target document
                target_doc = None
                for doc in word_app.Documents:
                    try:
                        if (doc.Name in target_doc_id or 
                            (hasattr(doc, 'FullName') and doc.FullName in target_doc_id)):
                            target_doc = doc
                            break
                    except Exception:
                        continue
                
                if not target_doc:
                    print(f"Target document not found: {target_doc_id}")
                    return False
                
                # Store original window state (but not selection)
                try:
                    original_window = word_app.ActiveWindow
                except Exception:
                    original_window = None
                
                try:
                    # Activate document
                    target_doc.Activate()
                    
                    if mode == PasteMode.END:
                        # Move to end of document
                        end_point = target_doc.Content.End - 1
                        target_doc.Range(end_point, end_point).Select()
                        
                        # Add spacing for content separation
                        selection = word_app.Selection
                        selection.InsertParagraphBefore()
                        selection.InsertParagraphBefore()
                        selection.Collapse(0)  # Collapse to end
                    
                    # Store the starting point of our paste operation
                    try:
                        start_point = word_app.Selection.Start
                    except Exception:
                        start_point = None
                    
                    # Paste with formatting preservation
                    word_app.Selection.PasteAndFormat(16)  # wdFormatOriginalFormatting
                    
                    # Important: Move cursor to end of pasted content
                    try:
                        if start_point is not None:
                            # Find the end of what we just pasted
                            current_selection = word_app.Selection
                            # Select from start of our paste to current position
                            pasted_range = target_doc.Range(start_point, current_selection.End)
                            # Move to the end of what we pasted
                            pasted_range.Collapse(0)  # 0 = Collapse to end
                            pasted_range.Select()
                    except Exception as e:
                        print(f"Warning: Could not position cursor after paste: {e}")
                    
                    return True
                    
                finally:
                    # Restore original window if needed
                    try:
                        if original_window:
                            original_window.Activate()
                    except Exception:
                        pass
                
        except Exception as e:
            print(f"Error pasting to active document: {str(e)}")
            return False

class DocumentSearcher:
    def __init__(self, folder_path: str, prefix_manager: PrefixManager):
        self.folder_path = Path(folder_path)
        self.document_index: Dict[str, DocumentInfo] = {}
        self.word_handler = WordHandler()
        self.min_token_length = 2  # Minimum length for search tokens
        self.prefix_manager = prefix_manager
    
    def index_documents(self):
        """Build comprehensive document index with enhanced metadata."""
        self.document_index.clear()
        try:
            for file_path in self.folder_path.glob("**/*.docx"):
                try:
                    # Get all file stats in a single system call
                    stats = file_path.stat()
                    
                    # Calculate relative path for folder display
                    rel_path = str(file_path.relative_to(self.folder_path).parent)
                    if rel_path == '.':
                        rel_path = ''
                    
                    # Create basic document info
                    doc_info = DocumentInfo(
                        path=file_path,
                        name=file_path.name,
                        last_modified=stats.st_mtime,
                        created_time=stats.st_ctime,
                        size=stats.st_size,
                        relative_path=rel_path
                    )
                    
                    # Try to read hierarchy metadata if it exists
                    try:
                        # First check for metadata JSON file
                        meta_file_path = file_path.with_suffix('.docx.meta.json')
                        if not meta_file_path.exists():
                            # Try alternative name format
                            meta_file_path = file_path.with_suffix('.meta.json')
                        
                        if meta_file_path.exists():
                            # Read metadata from JSON file
                            import json
                            with open(meta_file_path, 'r', encoding='utf-8') as f:
                                metadata = json.load(f)
                                
                            # Apply metadata to document info
                            if 'original_doc_path' in metadata:
                                doc_info.original_doc_path = metadata['original_doc_path']
                                
                            if 'position_in_original' in metadata:
                                doc_info.position_in_original = metadata['position_in_original']
                                
                            if 'parent_doc_name' in metadata:
                                doc_info.parent_doc_name = metadata['parent_doc_name']
                                
                            if 'sibling_docs' in metadata:
                                doc_info.sibling_docs = metadata['sibling_docs']
                            
                        # Fallback to document properties if no metadata file found
                        elif file_path.suffix.lower() == '.docx':
                            doc = docx.Document(file_path)
                            core_props = doc.core_properties
                            
                            # Get original document path
                            if hasattr(core_props, 'identifier') and core_props.identifier:
                                doc_info.original_doc_path = core_props.identifier
                            
                            # Get position in original document
                            if hasattr(core_props, 'category') and core_props.category and core_props.category.startswith('position:'):
                                try:
                                    doc_info.position_in_original = int(core_props.category.split(':')[1])
                                except (ValueError, IndexError):
                                    pass
                            
                            # Get parent document info
                            if hasattr(core_props, 'subject') and core_props.subject and core_props.subject.startswith('parent:'):
                                try:
                                    doc_info.parent_doc_name = core_props.subject.split(':')[1]
                                except IndexError:
                                    pass
                            
                            # Get sibling documents
                            if hasattr(core_props, 'comments') and core_props.comments:
                                try:
                                    import json
                                    doc_info.sibling_docs = json.loads(core_props.comments)
                                except (json.JSONDecodeError, ValueError):
                                    pass
                                    
                    except Exception as e:
                        print(f"Error reading document metadata for {file_path}: {e}")
                    
                    # Add to index
                    self.document_index[file_path.stem.lower()] = doc_info
                except (PermissionError, FileNotFoundError) as e:
                    print(f"Error accessing {file_path}: {e}")
                
        except Exception as e:
            print(f"Critical error during document indexing: {e}")
            raise

    def _extract_prefix(self, query: str) -> tuple[Optional[str], str]:
        """
        Extract prefix and remaining search terms from query.
        
        Args:
            query: Raw search query
            
        Returns:
            Tuple of (prefix if found, remaining search terms)
        """
        if not query:
            return None, ""
            
        # Split on first space
        parts = query.strip().split(maxsplit=1)
        if len(parts) != 2:
            return None, query.strip()
            
        potential_prefix, remainder = parts
        
        # Check if the first word is a valid prefix
        if self.prefix_manager.is_valid_prefix_word(potential_prefix):
            return potential_prefix, remainder.strip()
            
        return None, query.strip()

    def search(self, query: str, sort_key: str = None, reverse: bool = False, include_path: bool = False) -> List[DocumentInfo]:
        """
        Perform sophisticated multi-token search with optional sorting and prefix filtering.
        
        Args:
            query: Space-separated search terms, optionally starting with prefix
            sort_key: Optional key for sorting ('name', 'modified', 'created', 'size')
            reverse: Whether to reverse the sort order
            include_path: Whether to include path name in search text
            
        Returns:
            List of matching DocumentInfo objects, sorted if requested
        """
        # Handle empty query - now filters out excluded folders
        if not query.strip():
            results = []
            for doc in self.document_index.values():
                # Skip excluded folders when no prefix is used
                if not self.prefix_manager.is_folder_excluded(doc.relative_path):
                    results.append(doc)
            return self._sort_results(results, sort_key, reverse)

        # Extract prefix if present
        prefix, search_query = self._extract_prefix(query)
        
        # Get target folders if prefix specified
        target_folders = None
        if prefix:
            target_folders = self.prefix_manager.get_folders_for_prefix(prefix)
            if not target_folders:  # Fallback if prefix has no valid folders
                return []

        # Tokenize and filter search terms
        search_tokens = [
            token.lower() for token in search_query.split()
            if len(token) >= self.min_token_length
        ]
        
        # If all tokens are too short, return empty list
        if not search_tokens:
            return []

        # Perform search with folder filtering
        results = []
        for doc in self.document_index.values():
            # If using a prefix, only include documents from target folders
            if target_folders is not None:
                if not any(doc.relative_path.startswith(folder) for folder in target_folders):
                    continue
            # If no prefix is used, exclude documents from excluded folders
            elif self.prefix_manager.is_folder_excluded(doc.relative_path):
                continue
                
            # Determine what text to search in - either just filename or filename+path
            search_text = doc.search_name
            if include_path and doc.relative_path:
                search_text = f"{search_text} {doc.relative_path.lower()}"
                
            # Check if all tokens match
            if all(token in search_text for token in search_tokens):
                results.append(doc)

        return self._sort_results(results, sort_key, reverse)

    def get_document_context(self, doc_info: DocumentInfo) -> List[DocumentInfo]:
        """
        Get the document in its original context, with siblings in proper order.
        
        Args:
            doc_info: The document to get context for
            
        Returns:
            List of related documents in original document order
        """
        related_docs = []
        
        # If no hierarchy information, just return the document itself
        if not doc_info.original_doc_path and not doc_info.sibling_docs:
            return [doc_info]
            
        # Add current document
        related_docs.append(doc_info)
        
        # Add sibling documents if available
        for sibling_name in doc_info.sibling_docs:
            # Look for sibling in the index
            sibling_key = sibling_name.lower()
            if sibling_key in self.document_index:
                related_docs.append(self.document_index[sibling_key])
                
        # If we have a parent document, add it too
        if doc_info.parent_doc_name:
            parent_key = doc_info.parent_doc_name.lower()
            if parent_key in self.document_index:
                # Insert parent at the beginning
                related_docs.insert(0, self.document_index[parent_key])
                
        # Sort by position in original document if available
        # Group documents by levels (parent first, then siblings in order)
        # First, sort by whether it's a parent or not
        related_docs.sort(key=lambda d: 0 if d.parent_doc_name else 1)
        
        # Then sort siblings by position
        sorted_docs = []
        # First add parent documents (already at the beginning due to previous sort)
        parent_docs = [d for d in related_docs if not d.parent_doc_name]
        sorted_docs.extend(parent_docs)
        
        # Then add child documents sorted by position
        child_docs = [d for d in related_docs if d.parent_doc_name]
        child_docs.sort(key=lambda d: d.position_in_original if d.position_in_original is not None else float('inf'))
        sorted_docs.extend(child_docs)
        
        return sorted_docs
        
    def _sort_results(
        self,
        results: List[DocumentInfo],
        sort_key: Optional[str],
        reverse: bool
    ) -> List[DocumentInfo]:
        """
        Sort results based on specified criteria.
        
        Implements an extensible sorting system with support for multiple
        sort keys and proper handling of edge cases.
        """
        if not sort_key:
            return results

        # Define sorting key functions
        sort_functions = {
            'name': lambda x: x.name.lower(),
            'modified': lambda x: x.last_modified,
            'created': lambda x: x.created_time,
            'size': lambda x: x.size,
        }

        if sort_key not in sort_functions:
            print(f"Warning: Unknown sort key '{sort_key}', using default ordering")
            return results

        return sorted(
            results,
            key=sort_functions[sort_key],
            reverse=reverse
        )

class DocxSearchApp(QMainWindow):
    """
    Main application window with advanced document handling capabilities.
    
    This class implements a sophisticated document management interface that:
    - Provides real-time search capabilities with prefix-based folder filtering
    - Manages multiple document targeting modes
    - Implements intelligent state persistence
    - Handles complex user interactions with robust error recovery
    """

    hotkey_activated = pyqtSignal()

    def __init__(self):
        super().__init__()

        # First, try to connect to an existing instance
        socket = QLocalSocket()
        socket.connectToServer("DocxSearchAppLock")
        
        # If connection succeeds, another instance exists
        if socket.waitForConnected(500):  # 500ms timeout
            reply = QMessageBox.question(
                None,  # No parent since window isn't created yet
                "Instance Already Running",
                "Another instance of Block Search is already running.\n"
                "Would you like to start another instance?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.No:
                # Clean up and exit
                socket.disconnectFromServer()
                sys.exit(0)
            socket.disconnectFromServer()
        
        # Create server to prevent additional instances
        self.lock_server = QLocalServer(self)
        # Force remove any existing server in case of previous crash
        QLocalServer.removeServer("DocxSearchAppLock")
        # Start listening
        self.lock_server.listen("DocxSearchAppLock")
        
        self.setWindowTitle("Block Search")
        self.hotkey_activated.connect(self.handle_global_hotkey)
        self.setGeometry(100, 100, 800, 600)
        
        # Initialize persistent settings
        self.settings = QSettings('DocxSearchApp', 'Settings')
        
        # Load saved paths and states
        self.search_folder = self.settings.value('search_folder', os.getcwd())
        self.target_document = self.settings.value('target_document', None)
        self.active_target_id = None
        
        # Load shortcut configuration
        self.activation_shortcut = self.settings.value('activation_shortcut', 'ctrl+space')
        
        # Initialize prefix manager before document searcher
        self.prefix_manager = PrefixManager(self.settings)
        
        # Initialize core components with prefix manager
        self.searcher = DocumentSearcher(self.search_folder, self.prefix_manager)
        self.search_delay = 300  # ms for debouncing
        
        # Setup UI components
        self.setup_ui()
        self.setup_menu()
        
        # Initial document indexing
        self.index_documents()
        self._update_target_status()

        # Initialize system tray integration
        self.setup_system_tray()
        
        # Configure global hotkey
        self.setup_global_hotkey()
        
        # Don't show window initially
        self.show()

        # Initialize Win32 API constants
        self.SWP_NOMOVE = 0x0002
        self.SWP_NOSIZE = 0x0001
        self.SWP_NOZORDER = 0x0004
        self.SWP_FRAMECHANGED = 0x0020
        self.SWP_NOOWNERZORDER = 0x0200
        self.HWND_TOP = 0
        
        # Get handle for current window after it's created
        self.window_handle = None
        # Wait until window is properly initialized
        QTimer.singleShot(0, self._store_window_handle)

    def setup_system_tray(self):
        """Initialize system tray integration with sophisticated menu handling."""
        self.tray_icon = QSystemTrayIcon(self)
        
        # Create icon - you'll need to replace with your actual icon path
        icon = QIcon(r"C:\Mac\Home\Desktop\block_sender_icon.ico")
        self.tray_icon.setIcon(icon)
        
        # Create tray menu
        tray_menu = QTrayMenu()
        
        # Add menu actions
        show_action = tray_menu.addAction("Show Search")
        show_action.triggered.connect(self.activate_window)
        
        tray_menu.addSeparator()
        
        quit_action = tray_menu.addAction("Quit")
        quit_action.triggered.connect(self.quit_application)
        
        # Set the menu and make tray icon visible
        self.tray_icon.setContextMenu(tray_menu)
        self.tray_icon.show()
        
        # Connect double-click behavior
        self.tray_icon.activated.connect(self.on_tray_activated)
    
    def setup_global_hotkey(self):
        """Configure system-wide hotkey using saved shortcut configuration."""
        try:
            # First remove any existing hotkey
            keyboard.unhook_all()
            
            # Register new hotkey
            keyboard.add_hotkey(
                self.activation_shortcut,
                lambda: self.hotkey_activated.emit()
            )
            print(f"Hotkey for {self.activation_shortcut} has been registered successfully.")
        except Exception as e:
            print(f"Failed to register global hotkey: {e}")

    def handle_global_hotkey(self):
        """
        Handle global hotkey press with intelligent window management.
        Focuses window if open but not focused, otherwise toggles visibility.
        """
        if not self.isVisible():
            self.show()
            self.activate_window()
        elif not self.isActiveWindow():
            self.activate_window()
        else:
            self.focus_search()

    def configure_shortcut(self):
        """Show dialog for configuring the global activation shortcut."""
        dialog = ShortcutDialog(self.activation_shortcut, self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            new_shortcut = dialog.shortcut
            if new_shortcut != self.activation_shortcut:
                self.activation_shortcut = new_shortcut
                self.settings.setValue('activation_shortcut', new_shortcut)
                self.setup_global_hotkey()  # Reconfigure hotkey immediately

    def show_help(self):
        """Show the help dialog."""
        dialog = HelpDialog(self)
        dialog.exec()

    def show_document_splitter(self):
        """Show the document splitter dialog."""
        dialog = DocumentSplitterDialog(self)
        dialog.exec()
        
        # Refresh document index if the output is in the search folder
        if dialog.result() == QDialog.DialogCode.Accepted:
            output_dir = dialog.output_dir
            if output_dir and str(self.search_folder) in str(output_dir):
                self.index_documents()
                self.statusBar().showMessage("Document index refreshed with newly split documents", 3000)

    def _store_window_handle(self):
        """Store the Win32 window handle for later use."""
        self.window_handle = ctypes.windll.user32.GetParent(int(self.winId()))

    def _minimize_without_animation(self):
        """Minimize window without animation using direct Win32 API calls."""
        if self.window_handle:
            # Get current window placement info
            placement = wintypes.WINDOWPLACEMENT()
            placement.length = ctypes.sizeof(placement)
            ctypes.windll.user32.GetWindowPlacement(self.window_handle, ctypes.byref(placement))
            
            # Set minimized state directly
            placement.showCmd = 6  # SW_MINIMIZE
            ctypes.windll.user32.SetWindowPlacement(self.window_handle, ctypes.byref(placement))

    def _restore_without_animation(self):
        """Restore window without animation using direct Win32 API calls."""
        if self.window_handle:
            # Get current window placement info
            placement = wintypes.WINDOWPLACEMENT()
            placement.length = ctypes.sizeof(placement)
            ctypes.windll.user32.GetWindowPlacement(self.window_handle, ctypes.byref(placement))
            
            # Set restored state directly
            placement.showCmd = 1  # SW_NORMAL
            ctypes.windll.user32.SetWindowPlacement(self.window_handle, ctypes.byref(placement))
            
            # Ensure window is actually visible and focused
            flags = self.SWP_NOMOVE | self.SWP_NOSIZE | self.SWP_NOZORDER | \
                   self.SWP_FRAMECHANGED | self.SWP_NOOWNERZORDER
            ctypes.windll.user32.SetWindowPos(
                self.window_handle, 
                self.HWND_TOP, 
                0, 0, 0, 0, 
                flags
            )

    def activate_window(self):
        """
        Window activation with sophisticated Windows focus system integration.
        
        This implementation recognizes that Windows distinguishes between
        "explicit" and "ambient" focus states, and uses different strategies
        based on the current system focus state.
        """
        # First, store our current visibility state
        was_visible = self.isVisible()
        
        # If we're hidden or minimized, we get a "fresh" activation
        if not was_visible or (self.windowState() & Qt.WindowState.WindowMinimized):
            # Clear any existing window states and flags
            self.setWindowState(Qt.WindowState.WindowNoState)
            self.setWindowFlags(self.windowFlags() & ~Qt.WindowType.WindowStaysOnTopHint)
            
            # Show and activate - Windows is more likely to grant focus for newly shown windows
            self.show()
            self.activateWindow()
            self.search_input.setFocus()
        else:
            # For already-visible windows, we need to "release and reclaim" focus
            # This technique temporarily releases our claim on the window
            self.setWindowState(Qt.WindowState.WindowMinimized)
            
            # Use a very short timer to restore - this creates a natural focus transition
            QTimer.singleShot(10, self._restore_from_minimize)

    def _restore_from_minimize(self):
        """
        Restore window from minimized state using Windows-friendly focus claiming.
        
        The key insight is that restoring from a minimized state is treated
        differently by Windows' focus system compared to regular activation.
        """
        # Restore window state - this is a natural focus-claiming action
        self.setWindowState(Qt.WindowState.WindowNoState)
        self.show()
        self.activateWindow()
        
        # Now that we're restored, claim input focus
        self.search_input.setFocus()
        
        # Schedule a final focus check
        QTimer.singleShot(50, self._verify_focus)

    def _verify_focus(self):
        """
        Verify and rectify focus state if necessary.
        
        This final check ensures we achieved proper focus and provides
        a fallback if the main activation attempt failed.
        """
        if not self.search_input.hasFocus():
            self.activateWindow()
            self.search_input.setFocus()

    def focus_search(self):
        """Focus search field with intelligent text selection."""
        # Close context view if it's open
        if self.context_frame.isVisible():
            self.close_document_context()
            
        self.search_input.setFocus()
        self.search_input.selectAll()
    
    def on_tray_activated(self, reason):
        """Handle tray icon activation with platform-aware behavior."""
        if reason == QSystemTrayIcon.ActivationReason.DoubleClick:
            self.activate_window()

    def keyPressEvent(self, event: QKeyEvent):
        """Handle key press events at the main window level."""
        if event.key() == Qt.Key.Key_Escape:
            self.close_window()
            event.accept()
        else:
            super().keyPressEvent(event)

    def close_window(self):
        """Hide the window while keeping the application running."""
        self.hide()
    
    def quit_application(self):
        """Perform clean application shutdown with proper resource cleanup."""
        keyboard.unhook_all()
        self.tray_icon.hide()
        QApplication.quit()
    
    def closeEvent(self, event):
        """Override close event to implement minimize-to-tray behavior."""
        event.ignore()
        self.hide()

    def setup_ui(self):
        """Initialize the user interface with sophisticated component hierarchy."""
        # Create and configure central widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)
        
        # Target status panel with visual hierarchy
        target_frame = QFrame()
        target_frame.setFrameStyle(QFrame.Shape.StyledPanel | QFrame.Shadow.Sunken)
        # Set fixed height for target frame to prevent it from expanding too much
        target_frame.setFixedHeight(40)  # Fixed height in pixels
        target_layout = QHBoxLayout(target_frame)
        # Reduce margins to make the layout more compact
        target_layout.setContentsMargins(5, 3, 5, 3)
        
        target_label = QLabel("Target Document:")
        self.target_status = QLineEdit()
        self.target_status.setReadOnly(True)
        
        target_layout.addWidget(target_label)
        target_layout.addWidget(self.target_status)
        layout.addWidget(target_frame)
        
        # Search components with integrated keyboard navigation
        search_frame = QFrame()
        search_layout = QVBoxLayout(search_frame)
        
        # Create a horizontal splitter to hold search results and context panel
        self.search_splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # Initialize results list using our enhanced version
        self.results_list = EnhancedResultsList()
        self.results_list.setFocusPolicy(Qt.FocusPolicy.StrongFocus)
        # Connect signals
        self.results_list.itemActivated.connect(self.on_item_activated)
        self.results_list.ctrlEnterPressed.connect(self.on_ctrl_enter_activated)
        self.results_list.showDocumentContext.connect(self.show_document_context)
        self.results_list.focusSearchInput.connect(self.focus_search)
        
        # Create context panel (initially hidden)
        self.context_frame = QFrame()
        self.context_frame.setVisible(False)  # Hidden by default
        context_layout = QVBoxLayout(self.context_frame)
        
        # Create header for context view
        context_header = QFrame()
        context_header_layout = QHBoxLayout(context_header)
        context_header_layout.setContentsMargins(0, 0, 0, 0)
        
        self.context_title = QLabel("Document Context View")
        self.context_title.setStyleSheet("font-weight: bold; color: #2c5aa0;")
        context_header_layout.addWidget(self.context_title)
        
        context_header_layout.addStretch()
        
        # Add close button
        close_context_btn = QPushButton("×")  # Unicode × character
        close_context_btn.setToolTip("Close context view (Right Arrow)")
        close_context_btn.setMaximumWidth(30)
        close_context_btn.clicked.connect(self.close_document_context)
        context_header_layout.addWidget(close_context_btn)
        
        context_layout.addWidget(context_header)
        
        # Add document context list
        self.context_list = DocumentContextList()
        # Connect activate signal (Enter key)
        self.context_list.itemActivated.connect(self.on_context_item_activated)
        # Connect double-click signal explicitly
        self.context_list.itemDoubleClicked.connect(self.on_context_item_activated)
        # Connect other signals
        self.context_list.ctrlEnterPressed.connect(self.on_context_ctrl_enter_activated)
        self.context_list.closeContextView.connect(self.close_document_context)
        context_layout.addWidget(self.context_list)
        
        # Add preview shortcut (Shift+Enter) to context list
        context_preview_shortcut = QShortcut(QKeySequence(Qt.Key.Key_Return | Qt.KeyboardModifier.ShiftModifier), self.context_list)
        context_preview_shortcut.activated.connect(self.show_context_document_preview)
        
        # Add context panel and results list to splitter
        self.search_splitter.addWidget(self.context_frame)
        self.search_splitter.addWidget(self.results_list)
        
        # Set initial sizes (context panel takes less space)
        self.search_splitter.setSizes([0, 1])  # Context panel starts collapsed
        
        # Create enhanced search input with prefix support
        self.search_input = SearchInputWithKeyNavigation(
            self.results_list,
            self.prefix_manager
        )
        self.search_input.setPlaceholderText(
            f"Type to search documents in: {self.search_folder}"
        )
        self.search_input.textChanged.connect(self.on_search_text_changed)
        self.search_input.returnPressed.connect(self.handle_search_return)
        self.search_input.ctrlEnterPressed.connect(self.handle_ctrl_enter)
        
        # Add components to search layout
        search_layout.addWidget(self.search_input)
        search_layout.addWidget(self.search_splitter)
        layout.addWidget(search_frame)
        
        # Initialize status bar for user feedback
        self.statusBar().showMessage("Ready")
        
        # Configure search timer for debounced updates
        self.search_timer = QTimer()
        self.search_timer.setSingleShot(True)
        self.search_timer.timeout.connect(self.perform_search)

        # Preview shortcut (Shift+Enter)
        preview_shortcut = QShortcut(QKeySequence(Qt.Key.Key_Return | Qt.KeyboardModifier.ShiftModifier), self.results_list)
        preview_shortcut.activated.connect(self.show_document_preview)

    def handle_search_return(self):
        """Handle Return/Enter key in search box with intelligent focus management."""
        if self.results_list.count() > 0:
            current_item = self.results_list.currentItem()
            if current_item:
                # Use the CURRENT default mode
                self.on_item_activated(current_item)
            else:
                first_item = self.results_list.item(0)
                self.results_list.setCurrentItem(first_item)
                self.on_item_activated(first_item)

    def handle_ctrl_enter(self):
        """Handle Ctrl+Enter key in search box to use alternate paste mode."""
        print("Handle Ctrl+Enter called")
        if self.results_list.count() > 0:
            current_item = self.results_list.currentItem()
            if current_item:
                # Use the ALTERNATE mode
                self.on_ctrl_enter_activated(current_item)
            else:
                first_item = self.results_list.item(0)
                self.results_list.setCurrentItem(first_item)
                self.on_ctrl_enter_activated(first_item)

    def process_item_with_alternate_mode(self, item):
        """Process an item with the alternate paste mode."""
        # Get the file path from the item
        file_path = item.data(Qt.ItemDataRole.UserRole)
        success = False
        
        # Get the current paste mode and invert it
        use_cursor_mode = not self.cursor_mode_action.isChecked()
        
        if self.active_target_id:
            # Determine paste mode (inverted from default)
            paste_mode = PasteMode.CURSOR if use_cursor_mode else PasteMode.END
            
            # Paste to active document with the alternate mode
            success = self.searcher.word_handler.paste_to_active_document(
                file_path,
                self.active_target_id,
                mode=paste_mode
            )
            
            mode_str = "cursor position" if paste_mode == PasteMode.CURSOR else "document end"
            status_msg = (
                f"Pasted '{Path(file_path).name}' at {mode_str} (alternate mode)"
                if success
                else f"Error pasting to active document at {mode_str}"
            )
            
        elif self.target_document:
            # For target documents, we still use the standard transfer
            success = self.searcher.word_handler.transfer_content(
                file_path,
                self.target_document
            )
            status_msg = (
                f"Transferred '{Path(file_path).name}' to '{Path(self.target_document).name}'"
                if success
                else "Error transferring content. Check target document."
            )
            
        else:
            # Clipboard mode has no alternate
            success = self.searcher.word_handler.copy_to_clipboard(file_path)
            status_msg = (
                f"Copied '{Path(file_path).name}' to clipboard"
                if success
                else "Error copying to clipboard"
            )
        
        self.statusBar().showMessage(status_msg, 3000)

    def setup_menu(self):
        """Create application menus with comprehensive document management options."""
        menubar = self.menuBar()
        
        # Search settings menu
        settings_menu = menubar.addMenu('Search Settings')
        select_folder_action = settings_menu.addAction('Select Search Folder')
        select_folder_action.triggered.connect(self.select_search_folder)
        show_folder_action = settings_menu.addAction('Show Current Folder')
        show_folder_action.triggered.connect(self.show_current_folder)
        
        # Add Configure Shortcut option to Settings menu
        settings_menu.addSeparator()
        
        # Add option to include path names in search
        self.include_path_action = settings_menu.addAction('Include Path in Search')
        self.include_path_action.setCheckable(True)
        self.include_path_action.setChecked(self.settings.value('include_path_in_search', False, type=bool))
        self.include_path_action.setShortcut('Ctrl+Shift+P')  # Add keyboard shortcut
        self.include_path_action.triggered.connect(self.toggle_include_path)
        
        configure_shortcut_action = settings_menu.addAction('Set Window Focus Shortcut...')
        configure_shortcut_action.triggered.connect(self.configure_shortcut)
        
        # Add Prefix Configuration submenu
        prefix_menu = settings_menu.addMenu('Prefix Configuration')
        
        manage_prefixes_action = prefix_menu.addAction('Manage Prefixes...')
        manage_prefixes_action.triggered.connect(self.show_prefix_manager)
        
        prefix_menu.addSeparator()
        
        import_prefixes_action = prefix_menu.addAction('Import Prefixes...')
        import_prefixes_action.triggered.connect(self.import_prefixes)
        
        export_prefixes_action = prefix_menu.addAction('Export Prefixes...')
        export_prefixes_action.triggered.connect(self.export_prefixes)
        
        # Add Quit button to Settings menu
        settings_menu.addSeparator()
        quit_action = settings_menu.addAction('Quit')
        quit_action.setShortcuts([QKeySequence("Alt+F4"), QKeySequence("Ctrl+Q")])
        quit_action.triggered.connect(self.quit_application)

        # Document Tools menu
        doc_tools_menu = menubar.addMenu('Document Tools')
        split_doc_action = doc_tools_menu.addAction('Split Document by Headings...')
        split_doc_action.triggered.connect(self.show_document_splitter)

        # Target document menu
        target_menu = menubar.addMenu('Send to Closed Doc')
        select_target_action = target_menu.addAction('Select Destination...')
        select_target_action.setShortcut('Ctrl+T')
        select_target_action.triggered.connect(self.select_target_document)
        
        clear_target_action = target_menu.addAction('Clear Destination')
        clear_target_action.setShortcut('Ctrl+Shift+T')
        clear_target_action.triggered.connect(self.clear_target_document)
        
        target_menu.addSeparator()
        show_target_action = target_menu.addAction('Show Destination Info')
        show_target_action.triggered.connect(self.show_target_document)
        
        # Open Documents menu with paste mode options
        open_docs_menu = menubar.addMenu('Send to Open Doc')
        refresh_docs_action = open_docs_menu.addAction('Refresh Open Documents')
        refresh_docs_action.setShortcut('F5')
        refresh_docs_action.triggered.connect(self.refresh_open_documents)
        
        open_docs_menu.addSeparator()
        
        # Paste mode selection
        self.paste_mode_menu = menubar.addMenu('Default Paste Mode')

        # Create actions with dynamic shortcut display
        self.cursor_mode_action = QAction('Paste at Cursor', self)
        self.cursor_mode_action.setCheckable(True)
        self.cursor_mode_action.setChecked(True)
        
        self.end_mode_action = QAction('Paste at Document End', self)
        self.end_mode_action.setCheckable(True)
        
        # Create action group for mutual exclusivity
        paste_mode_group = QActionGroup(self)
        paste_mode_group.addAction(self.cursor_mode_action)
        paste_mode_group.addAction(self.end_mode_action)
        paste_mode_group.setExclusive(True)
        
        # Connect mode change to shortcut text update
        self.cursor_mode_action.triggered.connect(self._update_paste_mode_shortcuts)
        self.end_mode_action.triggered.connect(self._update_paste_mode_shortcuts)
        
        self.paste_mode_menu.addAction(self.cursor_mode_action)
        self.paste_mode_menu.addAction(self.end_mode_action)
        
        # Add a toggle shortcut action
        self.toggle_paste_mode_action = QAction('Toggle Paste Mode', self)
        self.toggle_paste_mode_action.setShortcut('Ctrl+P')
        self.toggle_paste_mode_action.triggered.connect(self.toggle_paste_mode)
        self.addAction(self.toggle_paste_mode_action)  # Add to application actions
        
        # Add the toggle action to the menu
        self.paste_mode_menu.addSeparator()
        self.paste_mode_menu.addAction(self.toggle_paste_mode_action)
        
        # Initialize shortcut displays
        self._update_paste_mode_shortcuts()
        
        # Store mode references
        open_docs_menu.addSeparator()
        self.open_docs_menu = open_docs_menu
        self.refresh_open_documents()

        # Add Sort menu
        sort_menu = menubar.addMenu('Sort')
        
        # Create sort options
        sort_group = QActionGroup(self)
        sort_group.setExclusive(True)
        
        # Define sort options with their display names and keys
        sort_options = [
            ('Name', 'name'),
            ('Date Modified', 'modified'),
            ('Date Created', 'created'),
            ('Size', 'size')
        ]
        
        # Create actions for each sort option
        self.sort_actions = {}
        for display_name, sort_key in sort_options:
            action = sort_menu.addAction(display_name)
            action.setCheckable(True)
            action.setData(sort_key)
            sort_group.addAction(action)
            self.sort_actions[sort_key] = action
            action.triggered.connect(self.perform_search)
        
        # Set default sort
        self.sort_actions['name'].setChecked(True)
        
        # Add reverse sort option
        sort_menu.addSeparator()
        self.reverse_sort_action = sort_menu.addAction('Reverse Order')
        self.reverse_sort_action.setCheckable(True)
        self.reverse_sort_action.triggered.connect(self.perform_search)
        
        # Add Help menu
        help_menu = menubar.addMenu('Help')
        show_help_action = help_menu.addAction('Show Help')
        show_help_action.setShortcut('F1')  # Standard help shortcut
        show_help_action.triggered.connect(self.show_help)

    def _update_paste_mode_shortcuts(self):
        """
        Update menu items to show appropriate shortcuts based on current default mode.
        """
        if self.cursor_mode_action.isChecked():
            self.cursor_mode_action.setText('Paste at Cursor (Enter)')
            self.end_mode_action.setText('Paste at Document End (Ctrl+Enter)')
        else:
            self.cursor_mode_action.setText('Paste at Cursor (Ctrl+Enter)')
            self.end_mode_action.setText('Paste at Document End (Enter)')

    def toggle_paste_mode(self):
        """Toggle between cursor and end paste modes."""
        if self.cursor_mode_action.isChecked():
            self.end_mode_action.setChecked(True)
        else:
            self.cursor_mode_action.setChecked(True)
        
        # Update the menu labels to reflect the new keyboard shortcuts
        self._update_paste_mode_shortcuts()
        
    def toggle_include_path(self):
        """Toggle whether path names are included in search terms."""
        include_path = self.include_path_action.isChecked()
        self.settings.setValue('include_path_in_search', include_path)
        
        # Show status message
        status = "enabled" if include_path else "disabled"
        self.statusBar().showMessage(f"Path name search {status} (Ctrl+Shift+P)", 3000)
        
        # Re-run search with new setting if search box has content
        if self.search_input.text():
            self.perform_search()

    def on_item_activated(self, item: QListWidgetItem):
        """Handle document selection using current default paste mode."""
        print("Regular item activation")
        file_path = item.data(Qt.ItemDataRole.UserRole)
        
        # Use the current default mode
        use_cursor_mode = self.cursor_mode_action.isChecked()
        self._process_item(file_path, use_cursor_mode)

    def on_ctrl_enter_activated(self, item: QListWidgetItem):
        """Handle document selection using alternate paste mode."""
        print("Ctrl+Enter item activation")
        file_path = item.data(Qt.ItemDataRole.UserRole)
        
        # Use the opposite of the current default mode
        use_cursor_mode = not self.cursor_mode_action.isChecked()
        self._process_item(file_path, use_cursor_mode)

    def show_document_preview(self, item=None):
        """Show preview of the selected document."""
        if item is None:
            item = self.results_list.currentItem()
            
        if not item:
            self.statusBar().showMessage("No document selected to preview", 3000)
            return
            
        file_path = item.data(Qt.ItemDataRole.UserRole)
        
        try:
            # Show preview dialog
            preview_dialog = DocumentPreviewDialog(file_path, self)
            preview_dialog.exec()
        except Exception as e:
            self.statusBar().showMessage(f"Error previewing document: {str(e)}", 5000)
            
    def show_context_document_preview(self):
        """Show preview of the selected document in context view."""
        item = self.context_list.currentItem()
            
        if not item:
            self.statusBar().showMessage("No document selected to preview", 3000)
            return
            
        file_path = item.data(Qt.ItemDataRole.UserRole)
        
        try:
            # Show preview dialog
            preview_dialog = DocumentPreviewDialog(file_path, self)
            preview_dialog.exec()
        except Exception as e:
            self.statusBar().showMessage(f"Error previewing document: {str(e)}", 5000)
            
    def show_document_context(self, doc_info: DocumentInfo):
        """Show document in the context of its original document structure."""
        # Check if document has context information
        related_docs = self.searcher.get_document_context(doc_info)
        
        # If no related documents, just return
        if len(related_docs) <= 1 and not doc_info.original_doc_path:
            self.statusBar().showMessage("No document context available", 3000)
            return
            
        # Clear the context list
        self.context_list.clear()
        
        # Get the index of the current document in the related docs list
        current_index = related_docs.index(doc_info) if doc_info in related_docs else 0
        
        # Set document context title
        if doc_info.original_doc_path:
            self.context_title.setText(f"Context: {Path(doc_info.original_doc_path).name}")
        else:
            self.context_title.setText("Document Context View")
        
        # Populate the context list
        for related_doc in related_docs:
            # Create item with document info
            item = QListWidgetItem()
            item.setText(related_doc.name)
            item.setData(Qt.ItemDataRole.UserRole, str(related_doc.path))
            item.setData(Qt.ItemDataRole.UserRole + 1, related_doc)  # Store doc_info
            
            # Add visual indicator if it's a parent document
            if related_doc.parent_doc_name:
                # Add visual prefix for child documents instead of indent
                item.setText("  ↪ " + item.text())  # Add arrow and spaces for child docs
            else:
                # Make parent docs bold
                font = item.font()
                font.setBold(True)
                item.setFont(font)
                
            self.context_list.addItem(item)
        
        # Select the current document in the context list
        if 0 <= current_index < self.context_list.count():
            self.context_list.setCurrentRow(current_index)
        
        # Apply visual effects to show context mode
        self.results_list.setEnabled(False)  # Disable the main results
        self.results_list.setStyleSheet("QListWidget { background-color: #f0f0f0; color: #808080; }")
        
        # Show context frame and give it focus
        self.context_frame.setVisible(True)
        
        # Update splitter sizes for better visibility
        total_width = self.search_splitter.width()
        self.search_splitter.setSizes([int(total_width * 0.4), int(total_width * 0.6)])
        
        # Give focus to the context list
        self.context_list.setFocus()
        
        # Update status bar
        self.statusBar().showMessage("Viewing document in context. Use arrow keys to navigate, Enter to select, Right arrow to exit", 5000)
    
    def close_document_context(self):
        """Close the document context view."""
        # Hide the context frame
        self.context_frame.setVisible(False)
        
        # Reset visual effects
        self.results_list.setEnabled(True)
        self.results_list.setStyleSheet("")
        
        # Update splitter sizes
        self.search_splitter.setSizes([0, self.search_splitter.width()])
        
        # Return focus to the main results list
        self.results_list.setFocus()
        
        # Update status bar
        self.statusBar().showMessage("Returned to search results", 3000)
    
    def on_context_item_activated(self, item: QListWidgetItem):
        """Handle item activation from the context list."""
        # Get document info
        path = item.data(Qt.ItemDataRole.UserRole)
        if not path:
            return
            
        # Use regular activation logic - same as on_item_activated
        use_cursor_mode = self.cursor_mode_action.isChecked()
        self._process_item(path, use_cursor_mode)
            
        # Keep context view open, just show status message
        self.statusBar().showMessage(f"Document sent to target. Use right arrow to exit context view.", 3000)
        
        # Return focus to the context list
        self.context_list.setFocus()
    
    def on_context_ctrl_enter_activated(self, item: QListWidgetItem):
        """Handle Ctrl+Enter activation from the context list."""
        # Get document info
        path = item.data(Qt.ItemDataRole.UserRole)
        if not path:
            return
            
        # Use alternate paste mode - same as on_ctrl_enter_activated
        use_cursor_mode = not self.cursor_mode_action.isChecked()
        self._process_item(path, use_cursor_mode)
            
        # Keep context view open, just show status message
        self.statusBar().showMessage(f"Document sent to target. Use right arrow to exit context view.", 3000)
        
        # Return focus to the context list
        self.context_list.setFocus()

    def _process_item(self, file_path: str, use_cursor_mode: bool):
        """Process a document with the specified paste mode."""
        success = False
        
        if self.active_target_id:
            # Set paste mode based on the provided mode flag
            paste_mode = PasteMode.CURSOR if use_cursor_mode else PasteMode.END
            
            # Paste to active document with selected mode
            success = self.searcher.word_handler.paste_to_active_document(
                file_path,
                self.active_target_id,
                mode=paste_mode
            )
            
            mode_str = "cursor position" if paste_mode == PasteMode.CURSOR else "document end"
            mode_type = " (default)" if use_cursor_mode == self.cursor_mode_action.isChecked() else " (alternate)"
            status_msg = (
                f"Pasted '{Path(file_path).name}' at {mode_str}{mode_type}"
                if success
                else f"Error pasting to active document at {mode_str}"
            )
            
        elif self.target_document:
            # Use existing file target logic
            success = self.searcher.word_handler.transfer_content(
                file_path,
                self.target_document
            )
            status_msg = (
                f"Transferred '{Path(file_path).name}' to '{Path(self.target_document).name}'"
                if success
                else "Error transferring content. Check target document."
            )
            
        else:
            # Clipboard mode
            success = self.searcher.word_handler.copy_to_clipboard(file_path)
            status_msg = (
                f"Copied '{Path(file_path).name}' to clipboard"
                if success
                else "Error copying to clipboard"
            )
        
        self.statusBar().showMessage(status_msg, 3000)
        
    def _update_target_status(self):
        """Update the target document status display with enhanced visibility."""
        if self.active_target_id:
            self.target_status.setText("Target: [Active Document]")
            self.target_status.setStyleSheet("""
                QLineEdit {
                    background-color: #e6f3ff;
                    color: #000000;
                    padding: 2px 5px;
                    border: 1px solid #b8d6f3;
                }
            """)
        elif self.target_document:
            target_name = Path(self.target_document).name
            self.target_status.setText(f"Target: {target_name}")
            self.target_status.setStyleSheet("""
                QLineEdit {
                    background-color: #e6ffe6;
                    color: #000000;
                    padding: 2px 5px;
                    border: 1px solid #b8f3b8;
                }
            """)
        else:
            self.target_status.clear()
            self.target_status.setPlaceholderText("No target document selected (using clipboard)")
            self.target_status.setStyleSheet("""
                QLineEdit {
                    background-color: #ffffff;
                    color: #666666;
                    padding: 2px 5px;
                    border: 1px solid #cccccc;
                }
                QLineEdit::placeholder {
                    color: #999999;
                }
            """)
    
    def refresh_open_documents(self):
        """Update the list of open documents in the menu."""
        # Clear existing document actions
        for action in self.open_docs_menu.actions()[2:]:  # Skip refresh and separator
            self.open_docs_menu.removeAction(action)
        
        # Get current open documents
        active_docs = self.searcher.word_handler.get_active_documents()
        
        if not active_docs:
            no_docs_action = self.open_docs_menu.addAction('No Open Documents')
            no_docs_action.setEnabled(False)
            return
        
        # Add action for each open document
        for doc in active_docs:
            action = self.open_docs_menu.addAction(doc.name)
            action.setData(doc.doc_id)
            action.triggered.connect(lambda checked, d=doc: self.set_active_target(d))
    
    def set_active_target(self, doc: ActiveDocument):
        """Set an open document as the paste target."""
        self.target_document = None  # Clear file target
        self.active_target_id = doc.doc_id
        self._update_target_status()
        self.statusBar().showMessage(f"Set active target to: {doc.name}", 3000)

        # Ensure screen updating is restored for the active document
        try:
            word_app = win32com.client.GetActiveObject("Word.Application")
            word_app.ScreenUpdating = True
        except:
            pass

    def select_target_document(self):
        """Open file dialog for selecting target document."""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Target Document",
            self.search_folder,
            "Word Documents (*.docx)"
        )
        
        if file_path:
            try:
                # Validate document
                with self.searcher.word_handler.word_session() as word_app:
                    with self.searcher.word_handler.open_document(word_app, file_path):
                        pass
                    word_app.ScreenUpdating = True

                self.active_target_id = None  # Clear active document target
                self.target_document = file_path
                self.settings.setValue('target_document', file_path)
                self._update_target_status()
                self.statusBar().showMessage(f"Target document set: {Path(file_path).name}", 3000)
                
            except Exception as e:
                QMessageBox.warning(
                    self,
                    "Invalid Target Document",
                    f"Could not use selected document as target: {str(e)}"
                )
    
    def clear_target_document(self):
        """Clear target document setting."""
        if self.target_document or self.active_target_id:
            reply = QMessageBox.question(
                self,
                "Clear Target Document",
                "Are you sure you want to clear the target document?\n"
                "Content will be copied to clipboard instead.",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                self.target_document = None
                self.active_target_id = None
                self.settings.setValue('target_document', None)
                self._update_target_status()
                self.statusBar().showMessage("Target document cleared", 3000)
    
    def show_target_document(self):
        """Display current target document information."""
        if self.active_target_id:
            QMessageBox.information(
                self,
                "Target Document",
                "Currently targeting an open document.\n"
                "Content will be pasted at the cursor position."
            )
        elif self.target_document:
            target_path = Path(self.target_document)
            try:
                modified_time = target_path.stat().st_mtime
                modified_str = datetime.fromtimestamp(modified_time).strftime('%Y-%m-%d %H:%M:%S')
                
                QMessageBox.information(
                    self,
                    "Target Document",
                    f"Current target document:\n"
                    f"Name: {target_path.name}\n"
                    f"Path: {target_path}\n"
                    f"Last modified: {modified_str}"
                )
            except Exception as e:
                QMessageBox.warning(
                    self,
                    "Target Document Error",
                    f"Error accessing target document:\n{str(e)}\n\n"
                    "You may want to clear and reselect the target document."
                )
        else:
            QMessageBox.information(
                self,
                "Target Document",
                "No target document set.\n"
                "Content will be copied to clipboard."
            )
    
    def index_documents(self):
        """Update document index and UI."""
        self.searcher.index_documents()
        self.statusBar().showMessage(f"Indexed {len(self.searcher.document_index)} documents")
    
    def select_search_folder(self):
        """Open folder selection dialog."""
        folder = QFileDialog.getExistingDirectory(
            self,
            "Select Folder to Search",
            self.search_folder,
            QFileDialog.Option.ShowDirsOnly
        )
        
        if folder:
            self.search_folder = folder
            self.settings.setValue('search_folder', folder)
            self.search_input.setPlaceholderText(f"Type to search documents in: {self.search_folder}")
            self.searcher = DocumentSearcher(self.search_folder, self.prefix_manager)
            self.index_documents()
            
            if self.search_input.text():
                self.perform_search()
    
    def show_current_folder(self):
        """Display current search folder information."""
        QMessageBox.information(
            self,
            "Current Search Folder",
            f"Currently searching in:\n{self.search_folder}\n\n"
            f"Number of indexed documents: {len(self.searcher.document_index)}"
        )
    
    def on_search_text_changed(self, text: str):
        """Handle search input changes with intelligent debouncing."""
        # Close context view if it's open
        if self.context_frame.isVisible():
            self.close_document_context()
            
        self.search_timer.stop()
        self.search_timer.start(self.search_delay)
    
    def perform_search(self):
        """Execute search operation with sophisticated result management."""
        query = self.search_input.text()
        
        # Store current selection state before updating
        current_selected_path = None
        current_item = self.results_list.currentItem()
        if current_item:
            current_selected_path = current_item.data(Qt.ItemDataRole.UserRole)
        
        # Cache focus state
        had_focus = self.results_list.hasFocus()
        
        # Clear results with minimal visual disruption
        self.results_list.clear()
        
        # Determine current sort configuration
        sort_key = None
        for action in self.sort_actions.values():
            if action.isChecked():
                sort_key = action.data()
                break
        
        reverse = self.reverse_sort_action.isChecked()
        
        # Get include_path setting
        include_path = self.include_path_action.isChecked()
        
        # Perform search with current parameters
        results = self.searcher.search(query, sort_key, reverse, include_path)
        
        # Efficiently populate results with intelligent selection management
        selection_restored = False
        first_item = None
        
        # Batch update for better performance
        self.results_list.setUpdatesEnabled(False)
        try:
            for doc in results:
                item = SearchResultItem(doc)  # Use our custom result item
                self.results_list.addItem(item)
                
                # Track first item for default selection
                if first_item is None:
                    first_item = item
                
                # Attempt to restore previous selection
                if current_selected_path and str(doc.path) == current_selected_path:
                    self.results_list.setCurrentItem(item)
                    selection_restored = True
        finally:
            self.results_list.setUpdatesEnabled(True)
        
        # Handle selection state
        if not selection_restored and first_item:
            self.results_list.setCurrentItem(first_item)
            self.results_list.scrollToItem(
                first_item,
                QListWidget.ScrollHint.PositionAtTop
            )
        
        # Restore focus if needed
        if had_focus:
            self.results_list.setFocus()
        
        # Update status with result count
        self.statusBar().showMessage(
            f"Found {len(results)} matching documents"
        )

    def on_item_activated(self, item: QListWidgetItem):
        """Handle document selection with intelligent content transfer."""
        file_path = item.data(Qt.ItemDataRole.UserRole)
        success = False
        
        if self.active_target_id:
            # Determine paste mode from menu state
            paste_mode = (PasteMode.CURSOR if self.cursor_mode_action.isChecked() 
                         else PasteMode.END)
            
            # Paste to active document with selected mode
            success = self.searcher.word_handler.paste_to_active_document(
                file_path,
                self.active_target_id,
                mode=paste_mode
            )
            
            mode_str = "cursor position" if paste_mode == PasteMode.CURSOR else "document end"
            status_msg = (
                f"Pasted '{Path(file_path).name}' at {mode_str}"
                if success
                else f"Error pasting to active document at {mode_str}"
            )
            
        elif self.target_document:
            # Use existing file target logic
            success = self.searcher.word_handler.transfer_content(
                file_path,
                self.target_document
            )
            status_msg = (
                f"Transferred '{Path(file_path).name}' to '{Path(self.target_document).name}'"
                if success
                else "Error transferring content. Check target document."
            )
            
        else:
            # Clipboard mode
            success = self.searcher.word_handler.copy_to_clipboard(file_path)
            status_msg = (
                f"Copied '{Path(file_path).name}' to clipboard"
                if success
                else "Error copying to clipboard"
            )
        
        self.statusBar().showMessage(status_msg, 3000)

    def show_prefix_manager(self):
        """Show the prefix configuration dialog."""
        dialog = PrefixManagerDialog(self.prefix_manager, self.search_folder, self)
        dialog.exec()
        
        # After dialog closes, verify all folders still exist
        missing = self.prefix_manager.verify_folders_exist(self.search_folder)
        if missing:
            message = "The following prefix mappings have missing folders:\n\n"
            for prefix, folder in missing:
                message += f"Prefix '{prefix}' -> folder '{folder}'\n"
            message += "\nPlease update or remove these mappings."
            
            QMessageBox.warning(
                self,
                "Missing Folders",
                message
            )
    
    def import_prefixes(self):
        """Import prefix configurations from CSV."""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Import Prefix Configuration",
            "",
            "CSV Files (*.csv)"
        )
        
        if file_path:
            if self.prefix_manager.import_from_csv(file_path):
                QMessageBox.information(
                    self,
                    "Import Successful",
                    "Prefix configurations have been imported successfully."
                )
            else:
                QMessageBox.warning(
                    self,
                    "Import Failed",
                    "Failed to import prefix configurations.\n"
                    "Please check the file format and try again."
                )
    
    def export_prefixes(self):
        """Export prefix configurations to CSV."""
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export Prefix Configuration",
            "",
            "CSV Files (*.csv)"
        )
        
        if file_path:
            if not file_path.lower().endswith('.csv'):
                file_path += '.csv'
                
            if self.prefix_manager.export_to_csv(file_path):
                QMessageBox.information(
                    self,
                    "Export Successful",
                    "Prefix configurations have been exported successfully."
                )
            else:
                QMessageBox.warning(
                    self,
                    "Export Failed",
                    "Failed to export prefix configurations.\n"
                    "Please check file permissions and try again."
                )

def main():
    """
    Application entry point with sophisticated initialization.
    
    This implementation provides:
    1. Proper system integration
    2. Resource management
    3. Error handling
    4. System tray support
    """
    try:
        app = QApplication(sys.argv)
        
        # Prevent application exit when last window closes
        app.setQuitOnLastWindowClosed(False)
        
        # Apply Fusion style for consistent cross-platform appearance
        app.setStyle('Fusion')
        
        window = DocxSearchApp()
        
        # Enter Qt event loop with proper system integration
        sys.exit(app.exec())
        
    except Exception as e:
        print(f"Critical error during application startup: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
